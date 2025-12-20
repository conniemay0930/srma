
# app.py
# =========================================================
# 一句話帶你完成 Meta-analysis（BYOK / Traditional Chinese）
# Author: Ya Hsin Yao
#
# 免責聲明：本工具僅供學術研究/教學用途，不構成醫療建議或法律意見；
# 使用者須自行驗證所有結果、數值、引用與全文內容；請勿上傳可識別之病人資訊。
#
# 校內資源/授權提醒（重要）：
# - 若文章來自校內訂閱（付費期刊/EZproxy/館藏系統），請勿將受版權保護之全文上傳至任何第三方服務或公開部署之網站。
# - 請遵守圖書館授權條款：避免大量下載/自動化批次擷取、避免共享全文給未授權者。
# - 若不確定是否可上傳：建議改用「本機版」或僅上傳你有權分享的開放取用全文（OA/PMC）。
#
# Privacy notice (BYOK):
# - Key only used for this session; do not use on untrusted deployments; do not upload identifiable patient info.
# =========================================================

from __future__ import annotations

import io
import re
import math
import json
import html
import time
from dataclasses import dataclass
from typing import Dict, List, Tuple, Optional, Any

import requests
import pandas as pd
import streamlit as st
import xml.etree.ElementTree as ET

# Optional: PDF text extraction (multiple backends; best-effort, no OCR by default)
try:
    from PyPDF2 import PdfReader  # type: ignore
    HAS_PYPDF2 = True
except Exception:
    HAS_PYPDF2 = False

try:
    import fitz  # PyMuPDF  # type: ignore
    HAS_FITZ = True
except Exception:
    HAS_FITZ = False

try:
    import pdfplumber  # type: ignore
    HAS_PDFPLUMBER = True
except Exception:
    HAS_PDFPLUMBER = False

# Optional: Plotly for forest plot
try:
    import plotly.graph_objects as go  # type: ignore
    HAS_PLOTLY = True
except Exception:
    HAS_PLOTLY = False

# Optional: Matplotlib fallback
try:
    from scipy.stats import chi2  # type: ignore
    HAS_SCIPY = True
except Exception:
    HAS_SCIPY = False

# Optional: Matplotlib fallback
try:
    import matplotlib.pyplot as plt  # type: ignore
    HAS_MPL = True
except Exception:
    HAS_MPL = False

# Optional: Word export
try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


# -------------------- Page config --------------------
st.set_page_config(page_title="一句話帶你完成 MA（繁體中文）", layout="wide")

# -------------------- Styles --------------------
CSS = """
<style>
:root{
  --bg: #ffffff;
  --muted: #6b7280;
  --line: #e5e7eb;
  --soft: #f7f7fb;
  --warn-bg:#fff7ed;
  --warn-line:#f59e0b;
  --ok-bg:#ecfdf5;
  --ok-line:#10b981;
  --bad-bg:#fef2f2;
  --bad-line:#ef4444;
}
.small { font-size: 0.92rem; color: var(--muted); }
.muted { color: var(--muted); }
.wrap { white-space: normal; }
.card { border: 1px solid var(--line); border-radius: 16px; padding: 0.95rem 1.05rem; background: var(--bg);
        margin-bottom: 0.9rem; box-shadow: 0 1px 0 rgba(0,0,0,0.03); }
.notice { border-left: 5px solid var(--warn-line); background: var(--warn-bg); padding: 0.95rem 1.05rem; border-radius: 14px; }
.kpi { border: 1px solid var(--line); border-radius: 16px; padding: 0.8rem 1rem; background: var(--soft); }
.kpi .label { font-size: 0.84rem; color: var(--muted); }
.kpi .value { font-size: 1.35rem; font-weight: 800; color: #111827; }
.badge { display:inline-block; padding:0.18rem 0.6rem; border-radius:999px; font-size:0.78rem; margin-right:0.35rem;
         border:1px solid rgba(0,0,0,0.06); background:#f3f4f6; }
.badge-ok { background: var(--ok-bg); border-color: rgba(16,185,129,0.25); color:#065f46; }
.badge-warn { background: #fef3c7; border-color: rgba(245,158,11,0.25); color:#92400e; }
.badge-bad { background: var(--bad-bg); border-color: rgba(239,68,68,0.25); color:#991b1b; }
.hr { border:none; border-top:1px solid #eef2f7; margin: 0.9rem 0; }
a { text-decoration: none; }
.red { color: #b91c1c; font-weight: 650; }
.green { color: #065f46; font-weight: 650; }
.flow { display:grid; grid-template-columns: 1fr; gap: 10px; }
.flow-row{ display:grid; grid-template-columns: 1fr; gap: 10px; }
.flow-box{ border:1px solid var(--line); border-radius: 14px; padding: 10px 12px; background: #fff; }
.flow-box .t{ font-weight: 800; margin-bottom: 2px; }
.flow-box .n{ color: var(--muted); font-size: 0.92rem; }
.flow-arrow{ text-align:center; color: var(--muted); font-size: 1.1rem; }
@media (min-width: 900px){ .flow-row{ grid-template-columns: 1fr 1fr; gap: 12px; } }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

# -------------------- Minimal i18n --------------------
I18N = {
    "zh-TW": {
        "app_title": "一句話帶你完成 MA",
        "author": "作者：Ya Hsin Yao",
        "lang_label": "介面語言",
        "tips_title": "小叮嚀／注意事項",
        "byok_title": "LLM（使用者自備 key）",
        "byok_toggle": "啟用 LLM（使用者自備 key）",
        "byok_notice": "Key only used for this session; do not use on untrusted deployments; do not upload identifiable patient info.",
        "byok_consent": "我了解並同意：不在不受信任的部署輸入 key、不上傳可識別病人資訊",
        "byok_clear": "Clear key",
        "settings": "設定",
        "links_title": "文獻與資源連結（可選）",
        "resolver": "OpenURL resolver base",
        "ezproxy": "EZproxy prefix",
        "search_settings": "檢索設定",
        "article_type": "文章類型 filter",
        "custom_filter": "自訂 filter（可留空）",
        "goal_mode": "研究目標取向",
        "question_notice": "輸入一句研究問題",
        "question_help": "例如：「FLACS 是否比傳統 phaco 好？」或「Diffractive vs nondiffractive EDOF IOL visual performance」。",
        "question_label": "Research question (one sentence)",
        "run": "開始（自動產出完整流程）",
        "tabs_overview": "總覽（PRISMA）",
        "tabs_step1": "Step 1 搜尋式（可手改）",
        "tabs_step2": "Step 2 可行性（SR/MA/NMA）",
        "tabs_step34": "Step 3+4 Screen by title & abstract（初篩）",
        "tabs_ft": "Step 4b Full text review",
        "tabs_extract": "Step 5 Data extraction（寬表）",
        "tabs_ma": "Step 6 MA + 森林圖",
        "tabs_rob2": "Step 6b ROB 2.0（含理由）",
        "tabs_ms": "Step 7 稿件草稿（分段呈現）",
        "tabs_diag": "Diagnostics",
        "pubmed_edit": "PubMed query（editable）",
        "pubmed_refetch": "以此搜尋式重新抓 PubMed",
        "pubmed_restore": "恢復為自動產生",
        "download_query": "下載搜尋式（txt）",
        "feas_title": "可行性掃描（既有 SR/MA/NMA）+ 綜合建議",
        "feas_optional": "（可選）BYOK：可行性報告",
        "records_none": "沒有抓到 records。",
        "ft_bulk_upload": "批次上傳 PDF（可選）",
        "ft_single_upload": "PDF 上傳（單篇，可選）",
        "ft_extract_text": "抽字（PDF→文字）",
        "ft_text_area": "Full-text text（可貼上；建議先 OCR 再貼/上傳）",
        "ft_ai_fill": "AI 讀全文 + 回填（對這篇）",
        "extract_schema": "Extraction schema（可自行規劃欄位；一行一欄）",
        "extract_quick_add": "快速新增一筆（一次輸入完再寫入寬表）",
        "extract_editor": "進階：一次編輯整張表（按儲存才 commit）",
        "extract_save": "儲存/commit 寬表修改",
        "ma_outcome_label": "Outcome label（手動輸入；用來篩選要做 MA 的列）",
        "ma_measure": "Effect measure（OR/RR/HR/MD/SMD）",
        "ma_run": "Run MA + 森林圖（按鈕執行）",
        "ms_generate": "（可選）BYOK：生成更完整稿件（請人工核對）",
        "export_docx": "匯出 Word（DOCX）",
        "pico_edit": "PICO/搜尋擴充（可手動修正）",
        "pico_apply": "套用修正後 PICO",
        "english_hint": "提醒：PubMed 搜尋式建議使用英文。若輸入中文，請在 Step 1 手動修正成英文關鍵字/MeSH（或啟用 LLM 自動翻譯）。",
    },
    "en": {
        "app_title": "From one question to Meta-analysis",
        "author": "Author: Ya Hsin Yao",
        "lang_label": "Interface language",
        "tips_title": "Notes / Warnings",
        "byok_title": "LLM (Bring Your Own Key)",
        "byok_toggle": "Enable LLM (BYOK)",
        "byok_notice": "Key only used for this session; do not use on untrusted deployments; do not upload identifiable patient info.",
        "byok_consent": "I understand and agree: do not use key on untrusted deployments; do not upload identifiable patient info",
        "byok_clear": "Clear key",
        "settings": "Settings",
        "links_title": "Institution links (optional)",
        "resolver": "OpenURL resolver base",
        "ezproxy": "EZproxy prefix",
        "search_settings": "Search settings",
        "article_type": "Article type filter",
        "custom_filter": "Custom filter (optional)",
        "goal_mode": "Goal mode",
        "question_notice": "Enter one research question",
        "question_help": "e.g., “FLACS vs conventional phaco” or “Diffractive vs nondiffractive EDOF IOL visual performance”.",
        "question_label": "Research question (one sentence)",
        "run": "Run (end-to-end pipeline)",
        "tabs_overview": "Overview (PRISMA)",
        "tabs_step1": "Step 1 Query (editable)",
        "tabs_step2": "Step 2 Feasibility (SR/MA/NMA)",
        "tabs_step34": "Step 3+4 Screen by title & abstract",
        "tabs_ft": "Step 4b Full-text review",
        "tabs_extract": "Step 5 Data extraction",
        "tabs_ma": "Step 6 MA + Forest",
        "tabs_rob2": "Step 6b RoB 2.0",
        "tabs_ms": "Step 7 Manuscript draft",
        "tabs_diag": "Diagnostics",
        "pubmed_edit": "PubMed query (editable)",
        "pubmed_refetch": "Refetch PubMed using this query",
        "pubmed_restore": "Restore auto query",
        "download_query": "Download query (txt)",
        "feas_title": "Feasibility scan + recommendations",
        "feas_optional": "(Optional) BYOK: feasibility report",
        "records_none": "No records retrieved.",
        "ft_bulk_upload": "Bulk upload PDFs (optional)",
        "ft_single_upload": "Upload PDF (optional)",
        "ft_extract_text": "Extract text (PDF → text)",
        "ft_text_area": "Full-text text (paste here; OCR recommended)",
        "ft_ai_fill": "AI full-text review + extraction (this record)",
        "extract_schema": "Extraction schema (one column per line)",
        "extract_quick_add": "Quick add (enter once, then append)",
        "extract_editor": "Advanced: edit table (commit on save)",
        "extract_save": "Save / commit edits",
        "ma_outcome_label": "Outcome label (manual input)",
        "ma_measure": "Effect measure (OR/RR/HR/MD/SMD)",
        "ma_run": "Run MA + forest plot",
        "ms_generate": "(Optional) BYOK: generate richer draft (verify manually)",
        "export_docx": "Export Word (DOCX)",
        "pico_edit": "PICO / expansions (editable)",
        "pico_apply": "Apply PICO edits",
        "english_hint": "Note: PubMed queries are usually best in English. If your question is non-English, edit Step 1 into English keywords/MeSH (or enable LLM auto-translation).",
    },
}

def t(key: str) -> str:
    lang = st.session_state.get("UI_LANG", "zh-TW")
    return I18N.get(lang, I18N["zh-TW"]).get(key, key)

# =========================================================
# Helpers
# =========================================================
def norm_text(x: Any) -> str:
    if x is None:
        return ""
    x = html.unescape(str(x))
    x = re.sub(r"\s+", " ", x).strip()
    return x

def short(s: str, n: int = 120) -> str:
    s = s or ""
    return (s[:n] + "…") if len(s) > n else s

def ensure_columns(df: Optional[pd.DataFrame], cols: List[str], default: Any = "") -> pd.DataFrame:
    if df is None or not isinstance(df, pd.DataFrame):
        return pd.DataFrame({c: [] for c in cols})
    for c in cols:
        if c not in df.columns:
            df[c] = default
    return df

def safe_float(x: Any) -> Optional[float]:
    try:
        if x is None:
            return None
        s = str(x).strip()
        if s == "":
            return None
        return float(s)
    except Exception:
        return None

def pretty_json(d: Any) -> str:
    try:
        return json.dumps(d, ensure_ascii=False, indent=2)
    except Exception:
        return str(d)

def json_from_text(s: str) -> Optional[dict]:
    if not s:
        return None
    s = s.strip()
    try:
        return json.loads(s)
    except Exception:
        pass
    m = re.search(r"\{.*\}", s, flags=re.S)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return None
    return None

def format_abstract(text: str) -> str:
    """Pretty-print PubMed abstracts.

    Goals:
    - Keep paragraphs readable (avoid '擠在一起').
    - If structured headings exist, normalize to:
      Purpose： / Methods： / Result： / Discussion：
    """
    t0 = (text or "").strip()
    if not t0:
        return ""

    # Normalize whitespace
    t0 = re.sub(r"\s*\n\s*", "\n", t0)
    t0 = re.sub(r"[ \t]+", " ", t0)

    # Insert section breaks for common structured headings
    # Accept both ':' and '：'
    headings = (
        "PURPOSE|OBJECTIVE|AIM|BACKGROUND|METHODS|MATERIALS AND METHODS|RESULTS|CONCLUSIONS|CONCLUSION|DISCUSSION"
    )
    t0 = re.sub(
        rf"(?<!\n)\b({headings})\s*[:：]\s*",
        r"\n\n\1: ",
        t0,
        flags=re.IGNORECASE,
    )

    # If still a single long paragraph, split on sentence boundaries
    if "\n" not in t0 and len(t0) > 800:
        t0 = re.sub(r"(?<=\.)\s+(?=[A-Z])", "\n\n", t0)

    # Normalize headings to the 4 labels requested (Traditional Chinese punctuation)
    # Map: Objective/Aim/Background -> Purpose；Conclusions -> Discussion (if no explicit Discussion)
    lines = []
    for block in t0.split("\n\n"):
        b = block.strip()
        if not b:
            continue
        m = re.match(r"^(?P<h>[A-Za-z ]+)\s*:\s*(?P<body>.*)$", b, flags=re.DOTALL)
        if m:
            h = m.group("h").strip().lower()
            body = m.group("body").strip()
            if h in ["purpose", "objective", "aim", "background"]:
                label = "Purpose："
            elif h in ["methods", "materials and methods"]:
                label = "Methods："
            elif h in ["results", "result"]:
                label = "Result："
            elif h in ["discussion", "conclusion", "conclusions"]:
                label = "Discussion："
            else:
                label = None

            if label:
                # Use bold label for readability in Streamlit markdown
                lines.append(f"**{label}** {body}" if body else f"**{label}**")
            else:
                lines.append(b)
        else:
            lines.append(b)

    out = "\n\n".join(lines).strip()

    # If we converted Conclusions to Discussion and also already have Discussion, keep both as separate paras
    return out

def badge(label: str) -> str:
    label = label or "Unsure"
    if label == "Include":
        return f"<span class='badge badge-ok'>{label}</span>"
    if label == "Exclude":
        return f"<span class='badge badge-bad'>{label}</span>"
    return f"<span class='badge badge-warn'>{label}</span>"

def to_csv_bytes(df: pd.DataFrame) -> bytes:
    if df is None:
        df = pd.DataFrame()
    return df.to_csv(index=False).encode("utf-8-sig")

def pubmed_link(pmid: str) -> str:
    pmid = str(pmid or "").strip().replace("PMID:", "").strip()
    return f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else ""

def doi_link(doi: str) -> str:
    doi = (doi or "").strip()
    return f"https://doi.org/{doi}" if doi else ""

def pmc_link(pmcid: str) -> str:
    pmcid = (pmcid or "").strip()
    if not pmcid:
        return ""
    if not pmcid.upper().startswith("PMC"):
        pmcid = "PMC" + pmcid
    return f"https://pmc.ncbi.nlm.nih.gov/articles/{pmcid}/"

# EZproxy / resolver utilities (no credentials stored)
def maybe_ezproxy(url: str) -> str:
    prefix = (st.session_state.get("EZPROXY_PREFIX") or "").strip()
    if not prefix:
        return url
    if not url:
        return url
    if url.startswith(prefix):
        return url
    return prefix + url

def resolver_url(doi: str, pmid: str = "") -> str:
    base = (st.session_state.get("RESOLVER_BASE") or "").strip()
    if not base:
        return ""
    doi = (doi or "").strip()
    if doi:
        return f"{base}sid=pmid:{pmid}&id=doi:{doi}"
    # fallback: at least provide PMID
    if pmid:
        return f"{base}sid=pmid:{pmid}"
    return ""

# =========================================================
# Protocol
# =========================================================
@dataclass
class Protocol:
    P: str = ""
    I: str = ""
    C: str = ""
    O: str = ""
    NOT: str = "animal OR mice OR rat OR in vitro OR case report"
    goal_mode: str = "Fast / feasible (gap-fill)"

    P_syn: List[str] = None
    I_syn: List[str] = None
    C_syn: List[str] = None
    O_syn: List[str] = None

    mesh_P: List[str] = None
    mesh_I: List[str] = None
    mesh_C: List[str] = None
    mesh_O: List[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "pico": {"P": self.P, "I": self.I, "C": self.C, "O": self.O, "NOT": self.NOT},
            "goal_mode": self.goal_mode,
            "search_expansion": {
                "P_synonyms": self.P_syn or [],
                "I_synonyms": self.I_syn or [],
                "C_synonyms": self.C_syn or [],
                "O_synonyms": self.O_syn or [],
                "NOT": [x.strip() for x in (self.NOT or "").split(" OR ") if x.strip()],
            },
            "mesh_candidates": {"P": self.mesh_P or [], "I": self.mesh_I or [], "C": self.mesh_C or [], "O": self.mesh_O or []},
        }

# =========================================================
# Session state
# =========================================================
def init_state():
    ss = st.session_state
    ss.setdefault("UI_LANG", "zh-TW")
    ss.setdefault("UI_LANG_DISPLAY", "繁體中文")

    # BYOK
    ss.setdefault("byok_enabled", False)
    ss.setdefault("byok_key", "")
    ss.setdefault("byok_base_url", "https://api.openai.com/v1")
    ss.setdefault("byok_model", "gpt-4o-mini")
    ss.setdefault("byok_temp", 0.2)
    ss.setdefault("byok_consent", False)

    ss.setdefault("byok_last_ok_endpoint", "")
    ss.setdefault("byok_last_error", "")

    # MeSH lookup (NLM MeSH RDF Lookup Service)
    ss.setdefault("mesh_lookup_enabled", True)
    ss.setdefault("mesh_lookup_limit", 6)
    ss.setdefault("mesh_lookup_match", "contains")


    # links
    ss.setdefault("RESOLVER_BASE", "")
    ss.setdefault("EZPROXY_PREFIX", "")

    # inputs
    ss.setdefault("question", "")
    ss.setdefault("article_type", "不限")
    ss.setdefault("custom_pubmed_filter", "")
    ss.setdefault("goal_mode", "Fast / feasible (gap-fill)")

    # PubMed fetch limits
    ss.setdefault("max_pubmed_records", 1000)
    ss.setdefault("pubmed_page_size", 200)

    # artifacts
    ss.setdefault("protocol", Protocol(P_syn=[], I_syn=[], C_syn=[], O_syn=[], mesh_P=[], mesh_I=[], mesh_C=[], mesh_O=[]))
    ss.setdefault("question_en", "")  # best-effort English
    ss.setdefault("pubmed_query_auto", "")
    ss.setdefault("pubmed_query", "")
    ss.setdefault("feas_query", "")
    ss.setdefault("pubmed_records", pd.DataFrame())
    ss.setdefault("srma_hits", pd.DataFrame())
    ss.setdefault("diagnostics", {})

    # TA screening
    ss.setdefault("ta_ai", {})
    ss.setdefault("ta_ai_reason", {})
    ss.setdefault("ta_ai_conf", {})
    ss.setdefault("ta_override", {})
    ss.setdefault("ta_override_reason", {})

    # Full text review
    ss.setdefault("ft_decision", {})   # pmid -> IncludeMA/Exclude/Not reviewed
    ss.setdefault("ft_reason", {})     # pmid -> reason string
    ss.setdefault("ft_pdf", {})        # pmid -> bytes
    ss.setdefault("ft_text", {})       # pmid -> str (extracted or pasted)

    # extraction
    ss.setdefault("extract_schema_text", "")
    ss.setdefault("extract_df", pd.DataFrame())
    ss.setdefault("extract_saved", False)

    # MA
    ss.setdefault("ma_outcome_input", "")
    ss.setdefault("ma_measure_choice", "OR")
    ss.setdefault("ma_model_choice", "Fixed effect")
    ss.setdefault("ma_last_result", None)
    ss.setdefault("ma_skipped_rows", pd.DataFrame())

    # ROB2
    ss.setdefault("rob2", {})

    # Manuscript
    ss.setdefault("ms_sections", {})
    ss.setdefault("writing_style_notes", "")

init_state()

# =========================================================
# BYOK LLM
# =========================================================
def llm_available() -> bool:
    return bool(st.session_state.get("byok_enabled")) and bool(st.session_state.get("byok_key", "").strip()) and bool(st.session_state.get("byok_consent"))

def call_openai_compatible(messages: List[Dict[str, str]], max_tokens: int = 1400) -> str:
    """Call an OpenAI-compatible endpoint.

    Strategy:
      1) Try Chat Completions: POST {base_url}/chat/completions
      2) If that fails, try Responses: POST {base_url}/responses

    This improves compatibility with providers/models that are moving to the Responses API.
    """
    base_url = (st.session_state.get("byok_base_url") or "").strip().rstrip("/")
    api_key = (st.session_state.get("byok_key") or "").strip()
    model = (st.session_state.get("byok_model") or "").strip()
    temperature = float(st.session_state.get("byok_temp") or 0.2)

    if not (base_url and api_key and model):
        raise RuntimeError("LLM 未設定完成（base_url / key / model）。")

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    def _record(ok_endpoint: str = "", err: str = ""):
        # Keep lightweight debug signals in UI.
        st.session_state["byok_last_ok_endpoint"] = ok_endpoint or ""
        st.session_state["byok_last_error"] = err or ""

    # ---- 1) Chat Completions ----
    try:
        url = f"{base_url}/chat/completions"
        payload = {"model": model, "messages": messages, "temperature": temperature, "max_tokens": max_tokens}
        r = requests.post(url, headers=headers, json=payload, timeout=75)
        if r.status_code == 200:
            data = r.json()
            out = (data.get("choices", [{}])[0].get("message", {}) or {}).get("content", "")
            if isinstance(out, str) and out.strip():
                _record(ok_endpoint="chat/completions", err="")
                return out
            # If response shape is odd, fall through to try /responses
            raise RuntimeError("chat/completions 回傳格式異常或空內容。")
        raise RuntimeError(f"chat/completions HTTP {r.status_code} / {r.text[:300]}")
    except Exception as e_chat:
        _record(ok_endpoint="", err=str(e_chat)[:300])

    # ---- 2) Responses API ----
    url = f"{base_url}/responses"
    payload = {
        "model": model,
        # Per OpenAI docs, simple message arrays are compatible as `input` in Responses.
        "input": messages,
        "temperature": temperature,
        "max_output_tokens": max_tokens,
    }
    r = requests.post(url, headers=headers, json=payload, timeout=75)
    if r.status_code != 200:
        _record(ok_endpoint="", err=f"responses HTTP {r.status_code} / {r.text[:300]}")
        raise RuntimeError(f"LLM 呼叫失敗：HTTP {r.status_code} / {r.text[:300]}")
    data = r.json()

    # Prefer official helper field when present.
    if isinstance(data, dict):
        if isinstance(data.get("output_text"), str) and data.get("output_text").strip():
            _record(ok_endpoint="responses", err="")
            return data["output_text"]

        # Otherwise, walk the typed `output` list.
        out_parts: List[str] = []
        for item in (data.get("output") or []):
            if not isinstance(item, dict):
                continue
            if item.get("type") != "message":
                continue
            for c in (item.get("content") or []):
                if not isinstance(c, dict):
                    continue
                if c.get("type") in ("output_text", "text"):
                    txt = c.get("text") or ""
                    if isinstance(txt, str) and txt:
                        out_parts.append(txt)
        out = "\\n".join([x for x in out_parts if x.strip()]).strip()
        if out:
            _record(ok_endpoint="responses", err="")
            return out

    _record(ok_endpoint="", err="responses 回傳格式異常或空內容。")
    raise RuntimeError("LLM 回傳格式異常或空內容。")


def llm_json(system: str, user: str, max_tokens: int = 1200) -> Optional[dict]:
    if not llm_available():
        return None
    try:
        out = call_openai_compatible(
            [{"role": "system", "content": system},
             {"role": "user", "content": user}],
            max_tokens=max_tokens,
        )
        return json_from_text(out)
    except Exception as e:
        st.warning(f"LLM 解析失敗：{e}")
        return None

# =========================================================
# Provider presets (OpenAI-compatible)
# =========================================================
PROVIDER_PRESETS = {
    # Official OpenAI API
    "OpenAI": {
        "base_url": "https://api.openai.com/v1",
        "models": [
            "gpt-4o-mini",
            "gpt-4o",
            "gpt-4.1-mini",
            "gpt-4.1",
            "o4-mini",
        ],
    },
    # OpenRouter (aggregator; requires OpenRouter key)
    "OpenRouter": {
        "base_url": "https://openrouter.ai/api/v1",
        "models": [
            "openai/gpt-4o-mini",
            "openai/gpt-4o",
            "anthropic/claude-3.5-sonnet",
            "google/gemini-1.5-pro",
            "meta-llama/llama-3.1-70b-instruct",
        ],
    },
    # Groq (fast; requires Groq key)
    "Groq": {
        "base_url": "https://api.groq.com/openai/v1",
        "models": [
            "llama-3.1-70b-versatile",
            "llama-3.1-8b-instant",
            "mixtral-8x7b-32768",
            "gemma2-9b-it",
        ],
    },
    # Local Ollama server (no external key; model names depend on what users pulled locally)
    "Ollama（本機）": {
        "base_url": "http://localhost:11434/v1",
        "models": [
            "llama3.1",
            "qwen2.5",
            "mistral",
        ],
    },
}

def _provider_model_options(provider_name: str) -> List[str]:
    if provider_name in PROVIDER_PRESETS:
        opts = PROVIDER_PRESETS[provider_name].get("models", []) or []
        # Always allow custom
        return list(opts) + ["（自訂）"]
    return ["（自訂）"]

# =========================================================
# Sidebar
# =========================================================
with st.sidebar:
    st.header(t("settings"))
    lang_choice = st.selectbox(
        t("lang_label"),
        options=["繁體中文", "English"],
        index=(0 if st.session_state.get("UI_LANG","zh-TW") == "zh-TW" else 1),
        key="UI_LANG_DISPLAY",
    )
    st.session_state["UI_LANG"] = ("zh-TW" if lang_choice == "繁體中文" else "en")

    with st.expander(t("tips_title"), expanded=False):
        st.markdown(
            "- **僅供學術用途、需人工核對**：輸出可能不完整或不正確，請逐一核對原文。\n"
            "- **勿上傳可識別病人資訊**：姓名、病歷號、影像、日期等。\n"
            "- **校內訂閱/授權全文上傳風險**：避免把校內訂閱 PDF 上傳到雲端（包含本 app 的雲端部署）；避免大量下載/批次擷取。\n"
            "- **PubMed/eUtils 被擋**：若抓不到文獻，請看 Diagnostics：esearch_url 是否回傳 HTML 或被防火牆擋；可改本機或換網路。\n"
            "- **未啟用 LLM 時自動降級**：不會卡在 extraction/ROB2；只提供模板/提示語，讓你人工補齊。\n"
            "- **PubMed query 建議英文**：中文問句會降低召回，請到 Step 1 改成英文關鍵字/MeSH，或啟用 LLM 自動翻譯。\n"
        )

    st.subheader(t("links_title"))
    st.text_input(t("resolver"), value=st.session_state.get("RESOLVER_BASE",""), key="RESOLVER_BASE",
                  help="填你們學校 link resolver base。每篇文獻會產生『學院全文連結』，使用者點了再自行登入下載。")
    st.text_input(t("ezproxy"), value=st.session_state.get("EZPROXY_PREFIX",""), key="EZPROXY_PREFIX",
                  help="若你們有 EZproxy，可填前綴。會把 DOI/Publisher/PMC 連結轉成 EZproxy 版本。")

    st.markdown("---")
    st.subheader(t("byok_title"))
    st.checkbox(t("byok_toggle"), key="byok_enabled")
    st.caption(t("byok_notice"))
    st.checkbox(t("byok_consent"), value=bool(st.session_state.get("byok_consent", False)), key="byok_consent")

    # Provider / Base URL / Model (OpenAI-compatible)
    provider = st.selectbox("Provider", options=list(PROVIDER_PRESETS.keys()) + ["自訂/相容API"], key="byok_provider")
    preset_base = PROVIDER_PRESETS.get(provider, {}).get("base_url", "https://api.openai.com/v1") if provider != "自訂/相容API" else st.session_state.get("byok_base_url","https://api.openai.com/v1")
    if provider != "自訂/相容API":
        st.caption(f"預設 Base URL：{preset_base}")
        if st.button("套用預設 Base URL", key="byok_apply_base"):
            st.session_state["byok_base_url"] = preset_base

    st.text_input("Base URL (OpenAI-compatible)", value=st.session_state.get("byok_base_url", preset_base), key="byok_base_url")

    model_opts = _provider_model_options(provider)
    current_model = st.session_state.get("byok_model", model_opts[0] if model_opts else "gpt-4o-mini")
    if current_model in model_opts:
        default_idx = model_opts.index(current_model)
    elif "（自訂）" in model_opts:
        default_idx = model_opts.index("（自訂）")
    else:
        default_idx = 0

    selected_model = st.selectbox("Model（下拉可選；也可自訂）", options=model_opts, index=default_idx, key="byok_model_dropdown")
    if selected_model == "（自訂）":
        st.text_input("自訂 Model ID", value=(current_model if current_model not in model_opts else ""), key="byok_model_custom", help="填完整 model id（依你選的 Provider/端點而定）。")
        st.session_state["byok_model"] = st.session_state.get("byok_model_custom","").strip() or current_model
    else:
        st.session_state["byok_model"] = selected_model

    st.text_input("API Key", type="password", key="byok_key")
    st.slider("Temperature", 0.0, 1.0, float(st.session_state.get("byok_temp", 0.2)), 0.05, key="byok_temp")
    # LLM debug status (helps identify silent fallback to heuristic mode)
    _ok_ep = (st.session_state.get("byok_last_ok_endpoint") or "").strip()
    _err = (st.session_state.get("byok_last_error") or "").strip()
    if _ok_ep:
        st.caption(f"LLM last OK endpoint: {_ok_ep}")
    if _err:
        st.caption(f"LLM last error: {_err}")

    st.button(t("byok_clear"), key="byok_clear_btn", on_click=lambda: st.session_state.update({"byok_key": ""}))

    st.markdown("---")
    st.subheader(t("search_settings"))
    # MeSH suggestion enhancement (no key required)
    st.checkbox("MeSH Lookup（NLM）", key="mesh_lookup_enabled",
                help="使用 NLM MeSH RDF Lookup Service 自動補全 MeSH Heading。LLM 掛掉時也能維持 MeSH 建議品質。")
    st.slider("MeSH 建議上限（每個詞）", 1, 15, int(st.session_state.get("mesh_lookup_limit", 6)), 1, key="mesh_lookup_limit")
    st.selectbox("MeSH Match 模式", options=["contains", "exact", "startswith"], key="mesh_lookup_match")

    st.selectbox(t("article_type"), options=["不限","RCT","SR/MA","NMA","Cohort","Case-control"], key="article_type")
    st.text_input(t("custom_filter"), key="custom_pubmed_filter", help="例如：humans[MeSH Terms] AND english[lang]；會 AND 到搜尋式內。")
    # PubMed fetch cap (default 1000); increase if you expect more records.
    _max_label = "最大抓取篇數（PubMed）" if st.session_state.get("UI_LANG","zh-TW") == "zh-TW" else "Max PubMed records to fetch"
    st.select_slider(_max_label, options=[200, 500, 1000, 2000, 5000], key="max_pubmed_records")
    st.caption("提示：抓取越多篇，越慢；PubMed/eUtils 有流量限制。若你預期 >5000，建議改用本機版或分批檢索。" if st.session_state.get("UI_LANG","zh-TW") == "zh-TW" else "Tip: higher limits run slower and may hit eUtils throttling. For >5000, use local/batched search.")
    st.selectbox(t("goal_mode"), options=["Fast / feasible (gap-fill)", "Rigorous / narrow scope"], key="goal_mode")

# =========================================================
# Header
# =========================================================
st.title(t("app_title"))
st.caption(f"{t('author')}　|　Language：{'繁體中文' if st.session_state.get('UI_LANG','zh-TW')=='zh-TW' else 'English'}　|　免責聲明：僅供學術用途；請自行驗證所有結果與引用。")

# =========================================================
# PICO parsing + expansions
# =========================================================
ABBR_MAP = {
    "EDOF": ["extended depth of focus", "extended depth-of-focus", "extended range of vision", "extended range-of-vision"],
    "IOL": ["intraocular lens", "intra-ocular lens"],
    "RCT": ["randomized controlled trial", "randomised controlled trial"],
    "NMA": ["network meta-analysis", "network meta analysis"],
    "FLACS": ["femtosecond laser-assisted cataract surgery", "femtosecond laser assisted cataract surgery"],
    "PHACO": ["phacoemulsification"],
}

ARTICLE_TYPE_FILTERS = {
    "不限": "",
    "RCT": "randomized controlled trial[pt] OR controlled clinical trial[pt] OR randomized[tiab] OR randomised[tiab]",
    "SR/MA": "systematic review[pt] OR meta-analysis[pt] OR \"systematic review\"[tiab] OR \"meta-analysis\"[tiab]",
    "NMA": "\"network meta-analysis\"[tiab] OR network meta analysis[tiab] OR NMA[tiab]",
    "Cohort": "cohort studies[MeSH Terms] OR cohort[tiab]",
    "Case-control": "case-control studies[MeSH Terms] OR case control[tiab]",
}

def has_cjk(s: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", s or ""))

def split_vs(question: str) -> Tuple[str, str]:
    q = question or ""
    m = re.split(r"\s+vs\.?\s+|\s+VS\.?\s+|\s+versus\s+| vs | VS ", q, flags=re.IGNORECASE)
    if len(m) >= 2:
        left = m[0].strip()
        right = " vs ".join([x.strip() for x in m[1:]]).strip()
        return left, right
    return q.strip(), ""

def expand_terms(text: str) -> List[str]:
    text = norm_text(text)
    if not text:
        return []
    syn: List[str] = []
    parts = re.split(r"[;,/]+", text)
    for p in parts:
        p = p.strip()
        if not p:
            continue
        syn.append(p)
        key = p.upper()
        if key in ABBR_MAP:
            syn.extend(ABBR_MAP[key])
        toks = re.findall(r"[A-Za-z]{2,18}", p)
        for t0 in toks:
            tu = t0.upper()
            if tu in ABBR_MAP:
                syn.extend(ABBR_MAP[tu])

    out, seen = [], set()
    for s in syn:
        s2 = s.strip()
        if not s2:
            continue
        k = s2.lower()
        if k not in seen:
            seen.add(k)
            out.append(s2)
    return out

def mesh_lookup_descriptors(label: str, limit: int = 6, match: str = "contains") -> List[str]:
    """Lookup MeSH descriptors by label (no API key needed).

    Uses NLM MeSH RDF Lookup Service:
      GET https://id.nlm.nih.gov/mesh/lookup/descriptor?label=...&match=...&limit=...
    """
    label = (label or "").strip()
    if not label:
        return []
    limit = int(limit or 6)
    limit = max(1, min(limit, 20))
    match = (match or "contains").strip() or "contains"

    url = "https://id.nlm.nih.gov/mesh/lookup/descriptor"
    params = {"label": label, "match": match, "limit": limit}
    try:
        r = requests.get(url, params=params, timeout=10)
        if r.status_code != 200:
            return []
        data = r.json()
        out = []
        for it in (data or []):
            if isinstance(it, dict) and (it.get("label") or "").strip():
                out.append(it["label"].strip())
        return out
    except Exception:
        return []


def propose_mesh_candidates(terms: List[str]) -> List[str]:
    """Propose MeSH headings for a list of terms.

    - Keeps your original heuristic mappings (fast, no network).
    - Optionally enhances via NLM MeSH RDF Lookup Service (if enabled).
    """
    mesh: List[str] = []

    # --- original heuristics (keep) ---
    for t0 in terms or []:
        tl = (t0 or "").lower()
        if "cataract" in tl:
            mesh += ["Cataract", "Cataract Extraction"]
        if "glaucoma" in tl:
            mesh += ["Glaucoma"]
        if "intraocular lens" in tl or "iol" in tl or "lens" in tl:
            mesh += ["Lenses, Intraocular", "Lens Implantation, Intraocular"]

    # --- NLM MeSH lookup (optional) ---
    if bool(st.session_state.get("mesh_lookup_enabled", True)):
        lim = int(st.session_state.get("mesh_lookup_limit", 6) or 6)
        match = (st.session_state.get("mesh_lookup_match", "contains") or "contains").strip()
        max_terms = 6
        for t0 in (terms or [])[:max_terms]:
            t0 = (t0 or "").strip()
            if not t0:
                continue
            mesh += mesh_lookup_descriptors(t0, limit=lim, match=match)

    # --- dedupe (case-insensitive) ---
    out, seen = [], set()
    for m0 in mesh:
        k = (m0 or "").strip().lower()
        if k and k not in seen:
            seen.add(k)
            out.append((m0 or "").strip())
    return out


def question_to_protocol_heuristic(question: str) -> Protocol:
    q = norm_text(question)
    left, right = split_vs(q)
    proto = Protocol(P="", I=left, C=right, O="", goal_mode=st.session_state.get("goal_mode","Fast / feasible (gap-fill)"))
    if proto.I and proto.C and proto.I.strip().lower() == proto.C.strip().lower():
        proto.C = "other comparator (different model/design)"
    proto.P_syn = expand_terms(proto.P)
    proto.I_syn = expand_terms(proto.I)
    proto.C_syn = expand_terms(proto.C)
    proto.O_syn = expand_terms(proto.O)
    proto.mesh_P = propose_mesh_candidates(proto.P_syn)
    proto.mesh_I = propose_mesh_candidates(proto.I_syn)
    proto.mesh_C = propose_mesh_candidates(proto.C_syn)
    proto.mesh_O = propose_mesh_candidates(proto.O_syn)
    return proto

def question_to_protocol_llm(question: str) -> Tuple[Protocol, str]:
    """
    Returns (protocol, question_en).

    - If LLM unavailable, fall back to heuristic and use original question as question_en.
    - LLM output is expected to be concept-level PICO (NOT PubMed syntax).
    """
    question = norm_text(question)
    if not llm_available():
        return question_to_protocol_heuristic(question), question

    sys = (
        "You are an evidence synthesis assistant. "
        "Given ONE research question (may be non-English), produce concept-level outputs (NOT PubMed query syntax): "
        "(1) translate to concise English (question_en); "
        "(2) propose PICO (P, I, C, O) in plain English; "
        "(3) propose synonym expansions for each axis (lists of short phrases/acronyms); "
        "(4) propose candidate MeSH descriptor labels (do not add [MeSH] tags). "
        "Return STRICT JSON only with keys: "
        "question_en, P, I, C, O, NOT, P_syn, I_syn, C_syn, O_syn, mesh_P, mesh_I, mesh_C, mesh_O. "
        "Rules: "
        "P/I/C/O must be short noun phrases (no boolean operators, no field tags). "
        "Synonym lists should be 3–8 items where applicable; include common abbreviations. "
        "MeSH lists should be 0–8 items where applicable. "
        "If you are uncertain, leave the field as an empty string or empty list. "
        "NOT should contain broad exclusions (e.g., animal, in vitro, case report) in plain text."
    )
    user = f"Question: {question}\\nReturn JSON only."
    d = llm_json(sys, user, max_tokens=900) or {}

    q_en = norm_text(d.get("question_en") or "")
    proto = Protocol(
        P=norm_text(d.get("P") or ""),
        I=norm_text(d.get("I") or ""),
        C=norm_text(d.get("C") or ""),
        O=norm_text(d.get("O") or ""),
        NOT=norm_text(d.get("NOT") or "animal OR mice OR rat OR in vitro OR case report"),
        goal_mode=st.session_state.get("goal_mode","Fast / feasible (gap-fill)"),
        P_syn=[norm_text(x) for x in (d.get("P_syn") or []) if norm_text(x)],
        I_syn=[norm_text(x) for x in (d.get("I_syn") or []) if norm_text(x)],
        C_syn=[norm_text(x) for x in (d.get("C_syn") or []) if norm_text(x)],
        O_syn=[norm_text(x) for x in (d.get("O_syn") or []) if norm_text(x)],
        mesh_P=[norm_text(x) for x in (d.get("mesh_P") or []) if norm_text(x)],
        mesh_I=[norm_text(x) for x in (d.get("mesh_I") or []) if norm_text(x)],
        mesh_C=[norm_text(x) for x in (d.get("mesh_C") or []) if norm_text(x)],
        mesh_O=[norm_text(x) for x in (d.get("mesh_O") or []) if norm_text(x)],
    )

    # Fallback: if any field blank, use heuristic fill
    if not (proto.I or proto.C or proto.P or proto.O):
        proto2 = question_to_protocol_heuristic(question)
        proto.P, proto.I, proto.C, proto.O = proto2.P, proto2.I, proto2.C, proto2.O

    if not q_en:
        q_en = question
    return proto, q_en

# =========================================================
# PubMed query builder
# =========================================================
def quote_tiab(term: str) -> str:
    term = term.strip()
    if not term:
        return ""
    if "[" in term and "]" in term:
        return term
    # if already quoted
    if term.startswith('"') and term.endswith('"'):
        return f"{term}[tiab]"
    return f"\"{term}\"[tiab]" if " " in term else f"{term}[tiab]"

def mesh_term(term: str) -> str:
    term = term.strip()
    if not term:
        return ""
    if "[" in term and "]" in term:
        return term
    return f"\"{term}\"[MeSH Terms]" if " " in term else f"{term}[MeSH Terms]"

def or_block(items: List[str]) -> str:
    items = [x.strip() for x in items if x and x.strip()]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return "(" + " OR ".join(items) + ")"

def build_pubmed_query(proto: Protocol, article_type: str, custom_filter: str) -> str:
    # Topic blocks: use synonyms (tiab) + mesh candidates
    def build_axis(text: str, syn: List[str], mesh: List[str]) -> str:
        tiabs = []
        if text.strip():
            tiabs.append(quote_tiab(text))
        for s in (syn or []):
            if s.strip() and s.strip().lower() != text.strip().lower():
                tiabs.append(quote_tiab(s))
        meshes = [mesh_term(m) for m in (mesh or []) if m.strip()]
        block = or_block([b for b in (or_block(meshes), or_block(tiabs)) if b])
        return block

    P_block = build_axis(proto.P, proto.P_syn or [], proto.mesh_P or [])
    I_block = build_axis(proto.I, proto.I_syn or [], proto.mesh_I or [])
    C_block = build_axis(proto.C, proto.C_syn or [], proto.mesh_C or [])
    O_block = build_axis(proto.O, proto.O_syn or [], proto.mesh_O or [])

    blocks = [b for b in [P_block, I_block, C_block, O_block] if b]
    if not blocks:
        blocks = [quote_tiab(norm_text(st.session_state.get("question_en") or st.session_state.get("question") or ""))]

    q = " AND ".join([f"({b})" if " OR " in b else b for b in blocks if b])

    # Article type filter (optional)
    f = ARTICLE_TYPE_FILTERS.get(article_type, "")
    if f.strip():
        q = f"({q}) AND ({f})"

    # Custom filter (optional)
    if (custom_filter or "").strip():
        q = f"({q}) AND ({custom_filter.strip()})"

    # NOT block
    NOT = (proto.NOT or "").strip()
    if NOT:
        q = f"({q}) NOT ({NOT})"

    return q

# =========================================================
# PubMed eUtils
# =========================================================
EUTILS = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"

def esearch(term: str, retstart: int = 0, retmax: int = 200) -> Tuple[List[str], int, str, List[str]]:
    warnings = []
    url = f"{EUTILS}/esearch.fcgi"
    params = {"db": "pubmed", "term": term, "retmode": "json", "retstart": retstart, "retmax": retmax}
    r = requests.get(url, params=params, timeout=45)
    esearch_url = r.url
    if r.status_code != 200:
        warnings.append(f"HTTP {r.status_code} from esearch")
        return [], 0, esearch_url, warnings
    # PubMed sometimes returns HTML if blocked
    if "application/json" not in (r.headers.get("content-type") or ""):
        warnings.append("esearch did not return JSON (maybe blocked by firewall / returned HTML).")
        return [], 0, esearch_url, warnings
    j = r.json()
    es = j.get("esearchresult", {})
    ids = es.get("idlist", []) or []
    try:
        count = int(es.get("count", 0))
    except Exception:
        count = 0
    return ids, count, esearch_url, warnings

def efetch(pmids: List[str]) -> Tuple[pd.DataFrame, List[str], List[str]]:
    warnings = []
    if not pmids:
        return pd.DataFrame(), [], warnings
    url = f"{EUTILS}/efetch.fcgi"
    params = {"db": "pubmed", "id": ",".join(pmids), "retmode": "xml"}
    r = requests.get(url, params=params, timeout=75)
    efetch_url = r.url
    if r.status_code != 200:
        warnings.append(f"HTTP {r.status_code} from efetch")
        return pd.DataFrame(), [efetch_url], warnings
    # Parse XML robustly
    try:
        root = ET.fromstring(r.text)
    except Exception:
        warnings.append("efetch XML parse failed (maybe blocked / truncated).")
        return pd.DataFrame(), [efetch_url], warnings

    rows = []
    for art in root.findall(".//PubmedArticle"):
        try:
            pmid = (art.findtext(".//PMID") or "").strip()
            title = norm_text(art.findtext(".//ArticleTitle") or "")
            abst_parts = [norm_text(x.text or "") for x in art.findall(".//Abstract/AbstractText")]
            abstract = "\n".join([p for p in abst_parts if p])
            journal = norm_text(art.findtext(".//Journal/Title") or "")
            year = norm_text(art.findtext(".//JournalIssue/PubDate/Year") or "") or norm_text(art.findtext(".//ArticleDate/Year") or "")
            pubtypes = [norm_text(x.text or "") for x in art.findall(".//PublicationType")]
            # Authors
            authors = []
            for a in art.findall(".//AuthorList/Author"):
                last = norm_text(a.findtext("LastName") or "")
                fore = norm_text(a.findtext("ForeName") or "")
                coll = norm_text(a.findtext("CollectiveName") or "")
                if coll:
                    authors.append(coll)
                else:
                    name = (fore + " " + last).strip()
                    if name:
                        authors.append(name)
            first_author = authors[0] if authors else ""
            # DOI / PMCID
            doi = ""
            pmcid = ""
            for idn in art.findall(".//ArticleIdList/ArticleId"):
                itype = (idn.get("IdType") or "").lower()
                val = norm_text(idn.text or "")
                if itype == "doi":
                    doi = val
                if itype == "pmc":
                    pmcid = val
            rows.append({
                "PMID": pmid,
                "Title": title,
                "Abstract": abstract,
                "Year": year,
                "Journal": journal,
                "First_author": first_author,
                "Authors": "; ".join(authors),
                "DOI": doi,
                "PMCID": pmcid,
                "PublicationTypes": "; ".join(pubtypes),
                "Source": "PubMed",
            })
        except Exception as e:
            warnings.append(f"Record parse error: {e}")
            continue

    df = pd.DataFrame(rows)
    return df, [efetch_url], warnings

def fetch_pubmed(term: str, max_records: int = 200, page_size: int = 200) -> Tuple[pd.DataFrame, Dict[str, Any]]:
    """
    Fetch PubMed records using ESearch pagination + EFetch batching.

    max_records: maximum number of records to retrieve (client-side cap)
    page_size: ESearch page size (<= 200 is usually safe)
    """
    diag: Dict[str, Any] = {"warnings": [], "esearch_urls": [], "efetch_urls": [], "pubmed_total_count": 0}
    term = (term or "").strip()
    if not term:
        diag["warnings"].append("Empty PubMed query.")
        return pd.DataFrame(), diag

    # First page to get count + first ids
    all_ids: List[str] = []
    retstart = 0
    remaining = int(max_records or 0)
    page_size = int(page_size or 200)
    if page_size <= 0:
        page_size = 200
    if page_size > 200:
        page_size = 200  # PubMed often caps practical retmax; keep safe

    # Iterate ESearch
    total_count = 0
    while True:
        this_retmax = min(page_size, max(0, remaining))
        if this_retmax <= 0:
            break
        ids, count, es_url, warn = esearch(term, retstart, this_retmax)
        diag["esearch_urls"].append(es_url)
        diag["warnings"].extend(warn)
        if retstart == 0:
            total_count = int(count or 0)
            diag["pubmed_total_count"] = total_count
        if not ids:
            break
        all_ids.extend([str(x) for x in ids])
        retstart += len(ids)
        remaining -= len(ids)

        # Stop if we reached the end per PubMed count or no progress
        if total_count and retstart >= total_count:
            break
        if len(ids) < this_retmax:
            break

        # light throttle
        time.sleep(0.34)

    # De-duplicate while preserving order
    seen = set()
    dedup_ids = []
    for pid in all_ids:
        if pid and pid not in seen:
            seen.add(pid)
            dedup_ids.append(pid)

    if not dedup_ids:
        return pd.DataFrame(), diag

    # EFetch in batches
    frames = []
    for i in range(0, len(dedup_ids), 200):
        batch = dedup_ids[i:i+200]
        df, ef_urls, warn2 = efetch(batch)
        diag["efetch_urls"].extend(ef_urls)
        diag["warnings"].extend(warn2)
        if not df.empty:
            frames.append(df)
        time.sleep(0.34)

    if not frames:
        return pd.DataFrame(), diag
    out = pd.concat(frames, ignore_index=True)

    # Ensure standard columns exist
    for c in ["PMID","Title","Abstract","Year","FirstAuthor","Journal","DOI","PubMedURL","PMC","PMCID"]:
        if c not in out.columns:
            out[c] = ""

    return out, diag


# =========================================================
# Screening heuristics (fallback when no LLM)
# =========================================================
def heuristic_ta(proto: Protocol, row: pd.Series) -> Tuple[str, str, float]:
    """
    Heuristic title/abstract screening.

    Goal: provide a defensible *audit trail* reason string that is more useful than raw keyword hits.
    """
    title = str(row.get("Title","") or "")
    abstract = str(row.get("Abstract","") or "")
    pubtypes = str(row.get("PublicationTypes","") or row.get("PublicationType","") or "")
    text = f"{title} {abstract}".strip()
    text_l = text.lower()
    pub_l = pubtypes.lower()

    # Hard exclusion cues
    if any(k in text_l for k in ["case report", "case series", "animal", "mice", "mouse", "rat", "in vitro"]):
        return "Exclude", "排除線索：非人體臨床研究（animal/in vitro/case report）。", 0.85

    # RCT / trial cues
    rct_pubtype = any(k in pub_l for k in ["randomized controlled trial", "controlled clinical trial"])
    trial_like = any(k in text_l for k in ["randomized", "randomised", "randomly", "trial", "controlled", "placebo", "double-blind", "single-blind"])
    is_rct = rct_pubtype or trial_like

    # PICO matching (lightweight)
    hits = []
    def _hit(axis_name: str, terms: List[str]) -> bool:
        for w in terms:
            w0 = (w or "").strip()
            if not w0:
                continue
            if w0.lower() in text_l:
                hits.append(f"{axis_name}：{w0}")
                return True
        return False

    _hit("I", (proto.I_syn or []) + ([proto.I] if proto.I else []))
    _hit("C", (proto.C_syn or []) + ([proto.C] if proto.C else []))
    _hit("P", (proto.P_syn or []) + ([proto.P] if proto.P else []))
    _hit("O", (proto.O_syn or []) + ([proto.O] if proto.O else []))

    # Article-type constraint (UI)
    want_rct = (st.session_state.get("article_type") == "RCT")
    if want_rct and not is_rct:
        strict = bool(st.session_state.get("strict_rct_screen", True))
        if strict:
            return "Exclude", "研究設計不符合：已選 RCT，但 PublicationTypes/摘要未見隨機對照試驗線索。", 0.82
        # non-strict: keep as Unsure if PICO strongly matches
        if len(hits) >= 2:
            return "Unsure", "PICO 可能符合，但未見 RCT 線索（PublicationTypes/摘要未標示隨機對照）。建議人工確認方法學。", 0.55
        return "Exclude", "研究設計不符合：已選 RCT，但 PublicationTypes/摘要未見隨機對照試驗線索。", 0.80

    # Decision thresholds
    if len(hits) >= 2:
        rs = "；".join(hits) if hits else "PICO 關鍵字符合。"
        if want_rct:
            rs += "；設計：RCT/試驗線索"
        return "Include", rs, 0.72 if not want_rct else 0.78

    if len(hits) == 1:
        rs = "；".join(hits) if hits else "部分符合"
        if want_rct and is_rct:
            rs += "；設計：RCT/試驗線索"
        return "Unsure", rs + "；資訊不足，保留至全文。", 0.55

    # No clear PICO hit; keep conservative unless explicitly filtered out by RCT rule above
    return "Unsure", "未見明確 PICO 命中；為避免漏納入，建議人工快速掃過摘要。", 0.45

# =========================================================
# Full-text utilities
# =========================================================
def extract_pdf_text(pdf_bytes: bytes, max_pages: int = 30) -> str:
    """Best-effort PDF text extraction (no OCR by default).

    Order: PyMuPDF (fitz) → pdfplumber → PyPDF2.
    """
    if not pdf_bytes:
        return ""

    max_pages = max(1, int(max_pages or 30))

    # 1) PyMuPDF: usually best for "real text" PDFs
    if 'HAS_FITZ' in globals() and HAS_FITZ:
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            parts = []
            for i in range(min(len(doc), max_pages)):
                try:
                    parts.append(doc.load_page(i).get_text("text") or "")
                except Exception:
                    continue
            text = "\n".join([p for p in parts if p]).strip()
            if len(text) >= 200:
                return text
        except Exception:
            pass

    # 2) pdfplumber: sometimes better on tricky layout
    if 'HAS_PDFPLUMBER' in globals() and HAS_PDFPLUMBER:
        try:
            parts = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    if i >= max_pages:
                        break
                    try:
                        parts.append(page.extract_text() or "")
                    except Exception:
                        continue
            text = "\n".join([p for p in parts if p]).strip()
            if len(text) >= 200:
                return text
        except Exception:
            pass

    # 3) PyPDF2 fallback
    if HAS_PYPDF2:
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            parts = []
            for i, page in enumerate(reader.pages):
                if i >= max_pages:
                    break
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    continue
            return "\n".join([p for p in parts if p]).strip()
        except Exception:
            return ""

    return 
def default_extraction_schema() -> str:
    # Starting point; user can edit schema text.
    # Includes both effect/CI fields and arm-level fields (so MD/SE/CI can be computed).
    return "\n".join([
        # IDs
        "StudyID",
        "PMID",
        "DOI",
        "First_author",
        "Year",
        "Journal",
        "Title",

        # Study characteristics
        "StudyDesign",
        "Setting",
        "Country",
        "RecruitmentPeriod",
        "FollowUp",
        "Funding",
        "Registration",

        # Eligibility / PICO
        "Population",
        "InclusionCriteria",
        "ExclusionCriteria",
        "Intervention",
        "Comparator",

        # Outcome
        "OutcomeLabel",
        "OutcomeUnit",
        "Timepoint",

        # Effect-size entry (preferred if reported)
        "Effect_measure",
        "Effect",
        "SE",
        "Lower_CI",
        "Upper_CI",

        # Dichotomous (optional)
        "Events_Treat",
        "Total_Treat",
        "Events_Control",
        "Total_Control",

        # Continuous (optional)
        "Mean_Treat",
        "SD_Treat",
        "N_Treat",
        "Mean_Control",
        "SD_Control",
        "N_Control",

        "Notes",
    ])
def build_data_extraction_prompt(proto: Protocol, schema_cols: List[str], full_text: str = "") -> str:
    schema = "\n".join([f"- {c}" for c in schema_cols if c.strip()])
    return f"""
You are helping with systematic review / meta-analysis data extraction.

Research question (English): {st.session_state.get('question_en','')}
PICO:
- P: {proto.P}
- I: {proto.I}
- C: {proto.C}
- O: {proto.O}

Key requirements:
1) If the PDF text looks OCR-noisy, still try; explicitly say what could not be read.
2) If tables/figures are referenced, look for them in the text; if not present, say "figure/table not available in extracted text".
3) Populate the extraction sheet columns (one row per study arm/outcome/timepoint as needed).
4) Do not invent numbers. If missing, leave blank.

Extraction sheet columns:
{schema}

Full text (may be truncated):
{(full_text[:12000] + '...') if full_text and len(full_text)>12000 else (full_text or '')}

Return:
- A short narrative summary of key outcomes
- Then a JSON array of rows, each row is an object with the exact column names above.
""".strip()

# =========================================================
# MA / Forest plot
# =========================================================
RATIO_MEASURES = {"OR","RR","HR"}

def se_from_ci(effect: float, lcl: float, ucl: float, measure: str) -> Optional[float]:
    if effect is None or lcl is None or ucl is None:
        return None
def _effect_to_theta_se(row: pd.Series, measure: str) -> Tuple[Optional[float], Optional[float], Optional[str]]:
    """Convert an effect with CI (or SE) into (theta, se) on the analysis scale.
    - For OR/RR/HR: analysis scale is log(effect)
    - For MD/SMD: analysis scale is effect itself
    """
    measure = (measure or "").upper().strip()
    eff = row.get("Effect")
    lcl = row.get("Lower_CI")
    ucl = row.get("Upper_CI")
    se0 = row.get("SE")

    try:
        if se0 is not None and str(se0).strip() != "":
            se_val = float(se0)
        else:
            se_val = None
    except Exception:
        se_val = None

    try:
        eff_val = float(eff)
        lcl_val = float(lcl) if lcl is not None and str(lcl).strip() != "" else None
        ucl_val = float(ucl) if ucl is not None and str(ucl).strip() != "" else None
    except Exception:
        return None, None, "Effect/CI not numeric"

    if measure in ["OR", "RR", "HR"]:
        if eff_val <= 0 or (lcl_val is not None and lcl_val <= 0) or (ucl_val is not None and ucl_val <= 0):
            return None, None, "Ratio measure must be > 0"
        theta = math.log(eff_val)
        if se_val is None:
            if lcl_val is None or ucl_val is None:
                return None, None, "Need CI or SE"
            se_val = (math.log(ucl_val) - math.log(lcl_val)) / (2 * 1.96)
        if se_val <= 0:
            return None, None, "SE invalid"
        return theta, se_val, None

    # MD / SMD
    theta = eff_val
    if se_val is None:
        if lcl_val is None or ucl_val is None:
            return None, None, "Need CI or SE"
        se_val = (ucl_val - lcl_val) / (2 * 1.96)
    if se_val <= 0:
        return None, None, "SE invalid"
    return theta, se_val, None


def _theta_to_effect(theta: float, se: float, measure: str) -> Tuple[float, float, float]:
    measure = (measure or "").upper().strip()
    if measure in ["OR", "RR", "HR"]:
        eff = math.exp(theta)
        lcl = math.exp(theta - 1.96 * se)
        ucl = math.exp(theta + 1.96 * se)
        return eff, lcl, ucl
    eff = theta
    lcl = theta - 1.96 * se
    ucl = theta + 1.96 * se
    return eff, lcl, ucl


def _heterogeneity(rows: List[Tuple[float, float]]) -> Tuple[float, float, float, float]:
    """Given list of (theta, var), compute Q, df, I2(%), tau2(DL)."""
    k = len(rows)
    if k <= 1:
        return 0.0, max(0, k - 1), 0.0, 0.0
    ws = [1.0 / v for _, v in rows]
    sumw = sum(ws)
    theta_fe = sum(w * th for (th, _), w in zip(rows, ws)) / sumw
    Q = sum(w * ((th - theta_fe) ** 2) for (th, _), w in zip(rows, ws))
    df = k - 1
    I2 = 0.0
    if Q > 0 and Q > df:
        I2 = max(0.0, (Q - df) / Q) * 100.0
    # DerSimonian-Laird tau2
    c = sumw - (sum(w * w for w in ws) / sumw if sumw > 0 else 0.0)
    tau2 = max(0.0, (Q - df) / c) if c > 0 else 0.0
    return Q, float(df), float(I2), float(tau2)


def fixed_effect_meta(df: pd.DataFrame, measure: str) -> Tuple[Optional[dict], pd.DataFrame]:
    """
    df must contain Effect + (Lower_CI, Upper_CI) or SE.
    Returns (result, skipped_df).
    """
    work = df.copy()
    work = ensure_columns(work, ["Effect","Lower_CI","Upper_CI","SE","StudyID","First_author","Year","Title","PMID"], default="")
    for c in ["Effect","Lower_CI","Upper_CI","SE"]:
        work[c] = pd.to_numeric(work[c], errors="coerce")

    rows = []
    skipped = []
    for _, r in work.iterrows():
        theta, se, err = _effect_to_theta_se(r, measure)
        if theta is None or se is None:
            rr = r.to_dict()
            rr["SkipReason"] = err or "Invalid"
            skipped.append(rr)
            continue
        var = se ** 2
        rows.append((theta, var, r.to_dict()))

    if not rows:
        return None, pd.DataFrame(skipped)

    # FE pooling
    ws = [1.0 / v for _, v, _ in rows]
    sumw = sum(ws)
    theta_hat = sum(w * th for (th, _, _), w in zip(rows, ws)) / sumw
    se_hat = math.sqrt(1.0 / sumw)
    pooled, pooled_lcl, pooled_ucl = _theta_to_effect(theta_hat, se_hat, measure)

    # Heterogeneity on theta scale
    Q, df_q, I2, tau2 = _heterogeneity([(th, v) for th, v, _ in rows])
    pQ = None
    if 'HAS_SCIPY' in globals() and HAS_SCIPY and df_q > 0:
        try:
            pQ = float(chi2.sf(Q, df_q))
        except Exception:
            pQ = None

    # per-study table on effect scale
    out = []
    for (th, v, rd), w in zip(rows, ws):
        se = math.sqrt(v)
        eff, lcl, ucl = _theta_to_effect(th, se, measure)
        out.append({
            "PMID": str(rd.get("PMID","")),
            "StudyID": str(rd.get("StudyID","")),
            "First_author": rd.get("First_author",""),
            "Year": rd.get("Year",""),
            "Title": rd.get("Title",""),
            "Effect": float(eff),
            "Lower_CI": float(lcl),
            "Upper_CI": float(ucl),
            "Weight": float(w),
        })
    tab = pd.DataFrame(out)
    tab["Weight_pct"] = tab["Weight"] / tab["Weight"].sum() * 100.0

    result = {
        "measure": measure.upper().strip(),
        "model": "Fixed effect",
        "k": len(rows),
        "pooled": pooled,
        "pooled_lcl": pooled_lcl,
        "pooled_ucl": pooled_ucl,
        "theta_hat": theta_hat,
        "se_hat": se_hat,
        "Q": Q,
        "df": df_q,
        "p_Q": pQ,
        "I2": I2,
        "tau2": tau2,
        "study_table": tab.sort_values("Weight", ascending=True).reset_index(drop=True),
    }
    return result, pd.DataFrame(skipped)


def random_effect_meta(df: pd.DataFrame, measure: str) -> Tuple[Optional[dict], pd.DataFrame]:
    """DerSimonian-Laird random-effects meta-analysis."""
    work = df.copy()
    work = ensure_columns(work, ["Effect","Lower_CI","Upper_CI","SE","StudyID","First_author","Year","Title","PMID"], default="")
    for c in ["Effect","Lower_CI","Upper_CI","SE"]:
        work[c] = pd.to_numeric(work[c], errors="coerce")

    rows = []
    skipped = []
    for _, r in work.iterrows():
        theta, se, err = _effect_to_theta_se(r, measure)
        if theta is None or se is None:
            rr = r.to_dict()
            rr["SkipReason"] = err or "Invalid"
            skipped.append(rr)
            continue
        var = se ** 2
        rows.append((theta, var, r.to_dict()))

    if not rows:
        return None, pd.DataFrame(skipped)

    Q, df_q, I2, tau2 = _heterogeneity([(th, v) for th, v, _ in rows])
    pQ = None
    if 'HAS_SCIPY' in globals() and HAS_SCIPY and df_q > 0:
        try:
            pQ = float(chi2.sf(Q, df_q))
        except Exception:
            pQ = None

    # RE pooling
    ws = [1.0 / (v + tau2) for _, v, _ in rows]
    sumw = sum(ws)
    theta_hat = sum(w * th for (th, _, _), w in zip(rows, ws)) / sumw
    se_hat = math.sqrt(1.0 / sumw)
    pooled, pooled_lcl, pooled_ucl = _theta_to_effect(theta_hat, se_hat, measure)

    out = []
    for (th, v, rd), w in zip(rows, ws):
        se = math.sqrt(v)
        eff, lcl, ucl = _theta_to_effect(th, se, measure)
        out.append({
            "PMID": str(rd.get("PMID","")),
            "StudyID": str(rd.get("StudyID","")),
            "First_author": rd.get("First_author",""),
            "Year": rd.get("Year",""),
            "Title": rd.get("Title",""),
            "Effect": float(eff),
            "Lower_CI": float(lcl),
            "Upper_CI": float(ucl),
            "Weight": float(w),
        })
    tab = pd.DataFrame(out)
    tab["Weight_pct"] = tab["Weight"] / tab["Weight"].sum() * 100.0

    result = {
        "measure": measure.upper().strip(),
        "model": "Random effects (DL)",
        "k": len(rows),
        "pooled": pooled,
        "pooled_lcl": pooled_lcl,
        "pooled_ucl": pooled_ucl,
        "theta_hat": theta_hat,
        "se_hat": se_hat,
        "Q": Q,
        "df": df_q,
        "p_Q": pQ,
        "I2": I2,
        "tau2": tau2,
        "study_table": tab.sort_values("Weight", ascending=True).reset_index(drop=True),
    }
    return result, pd.DataFrame(skipped)

def plot_forest(result: dict, title: str = ""):
    df = result.get("study_table")
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        st.info("沒有可畫的列。")
        return

    dfp = df.copy()
    dfp = ensure_columns(dfp, ["First_author","Year","Effect","Lower_CI","Upper_CI","Weight_pct","Title"], default="")
    dfp["StudyLabel"] = [
        f"{str(a).strip()} {str(y).strip()}".strip()
        for a, y in zip(dfp["First_author"].astype(str), dfp["Year"].astype(str))
    ]
    dfp = dfp.sort_values("Weight_pct", ascending=True).reset_index(drop=True)

    pooled = float(result.get("pooled"))
    pooled_lcl = float(result.get("pooled_lcl"))
    pooled_ucl = float(result.get("pooled_ucl"))
    measure = str(result.get("measure","")).upper().strip()
    model = str(result.get("model",""))

    ratio = measure in ["OR","RR","HR"]
    null = 1.0 if ratio else 0.0

    # axis limits
    xmin = float(dfp["Lower_CI"].min())
    xmax = float(dfp["Upper_CI"].max())
    xmin = min(xmin, pooled_lcl)
    xmax = max(xmax, pooled_ucl)
    if ratio:
        xmin = max(1e-4, xmin)
        xmax = max(xmin * 1.2, xmax)
        xmin = xmin / 1.3
        xmax = xmax * 1.3
    else:
        pad = (xmax - xmin) * 0.18 if xmax > xmin else 1.0
        xmin -= pad
        xmax += pad

    # layout
    k = dfp.shape[0]
    y = list(range(k))
    y_sum = k  # summary at bottom

    fig_h = max(3.6, 0.36 * (k + 3))
    fig, ax = plt.subplots(figsize=(9.6, fig_h))

    # plot CIs
    for i, row in dfp.iterrows():
        yi = y[i]
        eff = float(row["Effect"])
        lcl = float(row["Lower_CI"])
        ucl = float(row["Upper_CI"])
        wt = float(row.get("Weight_pct", 0.0))
        # CI line
        ax.plot([lcl, ucl], [yi, yi], lw=1.6)
        # square size proportional to weight
        ms = max(4.0, min(12.0, 4.0 + wt / 4.0))
        ax.plot([eff], [yi], marker="s", markersize=ms, linestyle="None")

    # pooled diamond
    dh = 0.28
    ax.fill([pooled_lcl, pooled, pooled_ucl, pooled], [y_sum, y_sum - dh, y_sum, y_sum + dh], alpha=0.35, edgecolor="black")
    ax.plot([pooled_lcl, pooled_ucl], [y_sum, y_sum], lw=1.0)

    # reference line
    ax.axvline(null, linestyle="--", lw=1.0)

    ax.set_ylim(-1, y_sum + 1)
    ax.set_yticks([])
    ax.set_xlim(xmin, xmax)
    if ratio:
        ax.set_xscale("log")
    ax.set_xlabel(f"{measure}")
    ax.set_title(title or f"Forest plot ({measure}, {model})")

    # text columns (RevMan-like)
    # Left: study labels; Right: effect (CI) and weight
    left_x = 0.02
    eff_x = 0.68
    wt_x = 0.92

    fig.text(left_x, 0.93, "Study or Subgroup", fontweight="bold")
    fig.text(eff_x, 0.93, f"{measure} (95% CI)", fontweight="bold")
    fig.text(wt_x, 0.93, "Weight", fontweight="bold", ha="right")

    # map y data to figure coordinates
    def _y_to_fig(yi: float) -> float:
        # convert axis data y to figure coordinate
        return ax.transData.transform((null, yi))[1] / fig.bbox.height

    for i, row in dfp.iterrows():
        yi = y[i]
        yf = _y_to_fig(yi)
        fig.text(left_x, yf, str(row["StudyLabel"])[:60])
        eff = float(row["Effect"]); lcl = float(row["Lower_CI"]); ucl = float(row["Upper_CI"])
        fig.text(eff_x, yf, f"{eff:.3g} [{lcl:.3g}, {ucl:.3g}]")
        fig.text(wt_x, yf, f"{float(row.get('Weight_pct',0.0)):.1f}%", ha="right")

    yf = _y_to_fig(y_sum)
    fig.text(left_x, yf, "Total (95% CI)", fontweight="bold")
    fig.text(eff_x, yf, f"{pooled:.3g} [{pooled_lcl:.3g}, {pooled_ucl:.3g}]", fontweight="bold")
    fig.text(wt_x, yf, "100.0%", fontweight="bold", ha="right")

    # heterogeneity footer
    Q = result.get("Q"); df_q = result.get("df"); pQ = result.get("p_Q"); I2 = result.get("I2"); tau2 = result.get("tau2")
    het = []
    if Q is not None and df_q is not None:
        if pQ is None:
            het.append(f"Chi²={Q:.2f}, df={int(df_q)}")
        else:
            het.append(f"Chi²={Q:.2f}, df={int(df_q)}, p={pQ:.3g}")
    if I2 is not None:
        het.append(f"I²={float(I2):.1f}%")
    if model.lower().startswith("random") and tau2 is not None:
        het.append(f"Tau²={float(tau2):.3g}")
    if het:
        fig.text(0.02, 0.03, "Heterogeneity: " + "; ".join(het))

    st.pyplot(fig, use_container_width=True)
def plot_rob2_traffic_light(rob2_map: Dict[str, dict], pmids: List[str], title: str = "RoB 2.0 traffic light"):
    """Simple RoB 2.0 traffic-light plot (RevMan-like)."""
    if not pmids:
        st.info("沒有可畫的研究。")
        return
    domains = list(ROB_DOMAINS) + ["Overall"]
    levels = {"Low": 0, "Some concerns": 1, "High": 2, "Unclear": 3}
    # Colors are not explicitly set elsewhere; for clarity here we set typical traffic-light colors.
    color_map = {
        "Low": "#2ecc71",
        "Some concerns": "#f1c40f",
        "High": "#e74c3c",
        "Unclear": "#bdc3c7",
    }

    data = []
    for pmid in pmids:
        r = (rob2_map or {}).get(pmid, {}) or {}
        row = []
        for d in domains:
            v = r.get(d, "Unclear")
            if v not in levels:
                v = "Unclear"
            row.append(v)
        data.append(row)

    k = len(pmids)
    fig_w = max(7.5, 0.75 * len(domains) + 2.0)
    fig_h = max(2.8, 0.36 * (k + 3))
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))

    ax.set_xlim(0, len(domains))
    ax.set_ylim(0, k)
    ax.invert_yaxis()

    for i in range(k):
        for j, d in enumerate(domains):
            v = data[i][j]
            rect = plt.Rectangle((j, i), 1, 1, facecolor=color_map.get(v, "#bdc3c7"), edgecolor="white", lw=1)
            ax.add_patch(rect)

    ax.set_xticks([j + 0.5 for j in range(len(domains))])
    ax.set_xticklabels(domains, rotation=30, ha="right")
    ax.set_yticks([i + 0.5 for i in range(k)])
    ax.set_yticklabels([str(p) for p in pmids])
    ax.set_title(title)
    ax.set_frame_on(False)

    # Legend
    handles = [plt.Rectangle((0, 0), 1, 1, color=color_map[k]) for k in ["Low", "Some concerns", "High", "Unclear"]]
    ax.legend(handles, ["Low", "Some concerns", "High", "Unclear"], loc="upper right", frameon=False)

    st.pyplot(fig, use_container_width=True)

def generate_manuscript_basic(proto: Protocol, prisma: dict, ma_result: Optional[dict]) -> Dict[str, str]:
    """IMRaD baseline draft with placeholders 『』 for missing items."""
    topic = st.session_state.get("question_en") or st.session_state.get("question") or ""
    topic = topic.strip() if topic else "『research question』"

    # PRISMA counts
    n_id = prisma.get("identified", "『』")
    n_scr = prisma.get("screened", "『』")
    n_ft = prisma.get("fulltext_assessed", "『』")
    n_inc = prisma.get("included", "『』")

    # MA stats
    ma_txt = "『meta-analysis summary』"
    if ma_result:
        pooled = ma_result.get("pooled")
        lcl = ma_result.get("pooled_lcl")
        ucl = ma_result.get("pooled_ucl")
        k = ma_result.get("k", 0)
        model = ma_result.get("model","Fixed effect")
        measure = ma_result.get("measure","")
        I2 = ma_result.get("I2", 0)
        try:
            ma_txt = f"{model} pooled {measure}={float(pooled):.3g} (95% CI {float(lcl):.3g}–{float(ucl):.3g}); k={k}; I²={float(I2):.1f}%"
        except Exception:
            pass

    intro = f"""Introduction
Rationale: Evidence on {topic} remains heterogeneous and requires quantitative synthesis to inform practice and future research.
Objective: To systematically review and meta-analyze comparative studies addressing {topic}.
""".strip()

    methods = f"""Methods
Protocol and reporting: This review followed PRISMA guidance. The protocol specified eligibility criteria, outcomes, and analytic methods. 『protocol registration/ID if applicable』
Eligibility criteria: Population={proto.P or '『』'}; Intervention={proto.I or '『』'}; Comparator={proto.C or '『』'}; Outcomes={proto.O or '『』'}. Study design constraints: {st.session_state.get('article_type','不限')}.
Information sources and search strategy: We searched PubMed using a strategy combining MeSH and free-text terms (see Step 1 for the full query). Additional sources included reference lists and related-citation screening as feasible. 『date of last search』
Study selection: Two reviewers (『names/roles』) screened titles/abstracts, followed by full-text assessment; disagreements were resolved by consensus/third reviewer. Screening counts: identified={n_id}, screened={n_scr}, full-text assessed={n_ft}, included={n_inc}.
Data extraction: We extracted study characteristics, PICO details, and effect estimates (with variance). When needed, effect sizes were derived from arm-level data. 『software/tools』
Risk of bias: We assessed risk of bias using RoB 2.0 for randomized trials, including domain-level judgments and overall risk.
Synthesis: Meta-analyses used inverse-variance weighting with fixed-effect and/or random-effects (DerSimonian–Laird) models as specified. Heterogeneity was quantified using Chi² (Q) and I² statistics.
""".strip()

    results = f"""Results
Study selection: We identified {n_id} records. After screening {n_scr} titles/abstracts, {n_ft} full texts were assessed and {n_inc} studies were included in quantitative synthesis.
Study characteristics: Included studies were published in 『years range』 and evaluated {proto.I or '『intervention』'} versus {proto.C or '『comparator』'} in {proto.P or '『population』'}. 『add key baseline characteristics and follow-up』
Risk of bias: Overall RoB 2.0 judgments were 『summary of Low/Some concerns/High』 with common issues in 『domains』.
Quantitative synthesis: {ma_txt}. 『subgroup/sensitivity analyses if applicable』
""".strip()

    discussion = f"""Discussion
Principal findings: This synthesis suggests that {proto.I or 'the intervention'} compared with {proto.C or 'the comparator'} is associated with 『direction/magnitude』 for {proto.O or 'the outcome'}.
Interpretation: Findings should be interpreted considering heterogeneity (I²) and risk of bias. 『clinical relevance』
Limitations: Potential limitations include database restriction to PubMed, incomplete reporting in primary studies, and residual heterogeneity. 『publication bias/short follow-up』
Implications: Future trials should standardize outcomes and report arm-level data to enable robust meta-analysis.
""".strip()

    return {
        "Introduction": intro,
        "Methods": methods,
        "Results": results,
        "Discussion": discussion,
    }
def manuscript_llm_enhance(proto: Protocol, prisma: dict, ma_result: Optional[dict], style_notes: str) -> Optional[Dict[str, str]]:
    if not llm_available():
        return None
    sys = (
        "You are a senior medical writer. Draft a high-quality meta-analysis manuscript in academic ENGLISH, "
        "IMRaD style, with clear paragraphs and cautious claims. "
        "You MUST NOT fabricate data. If information is missing, use placeholders wrapped in 『』. "
        "Return STRICT JSON with keys: Introduction, Methods, Results, Discussion."
    )
    user = {
        "topic_question": st.session_state.get("question_en") or st.session_state.get("question") or "",
        "PICO": proto.to_dict().get("pico"),
        "PRISMA_counts": prisma,
        "MA_summary": (None if not ma_result else {k: ma_result[k] for k in ["measure","model","k","pooled","pooled_lcl","pooled_ucl"]}),
        "style_notes": style_notes or "",
    }
    d = llm_json(sys, json.dumps(user, ensure_ascii=False), max_tokens=1600)
    if not d:
        return None
    out = {}
    for k in ["Introduction","Methods","Results","Discussion"]:
        out[k] = str(d.get(k) or "").strip()
    return out

# =========================================================
# PRISMA counts
# =========================================================
def compute_prisma(records: pd.DataFrame) -> dict:
    records = records if isinstance(records, pd.DataFrame) else pd.DataFrame()
    n_identified = int(records.shape[0]) if not records.empty else 0

    ta = st.session_state.get("ta_override") or {}
    # default to AI if no override
    def ta_label(pmid: str) -> str:
        if pmid in ta and ta[pmid]:
            return ta[pmid]
        return st.session_state.get("ta_ai", {}).get(pmid, "Unsure")

    screened = n_identified
    ta_excluded = 0
    ta_kept = 0
    for pmid in (records["PMID"].astype(str).tolist() if not records.empty and "PMID" in records.columns else []):
        lbl = ta_label(pmid)
        if lbl == "Exclude":
            ta_excluded += 1
        else:
            ta_kept += 1

    # Full-text decisions
    ft_dec = st.session_state.get("ft_decision") or {}
    ft_assessed = 0
    ft_excluded = 0
    ft_included_ma = 0
    reasons = st.session_state.get("ft_reason") or {}
    reason_counts = {}
    for pmid, d in ft_dec.items():
        if d and d != "Not reviewed":
            ft_assessed += 1
        if d == "Exclude":
            ft_excluded += 1
            rs = reasons.get(pmid,"Unspecified")
            reason_counts[rs] = reason_counts.get(rs, 0) + 1
        if d == "Include for meta-analysis":
            ft_included_ma += 1

    return {
        "records_identified": n_identified,
        "records_screened_ta": screened,
        "records_excluded_ta": ta_excluded,
        "records_for_full_text": ta_kept,
        "full_text_assessed": ft_assessed,
        "full_text_excluded": ft_excluded,
        "full_text_exclusion_reasons": reason_counts,
        "studies_included_ma": ft_included_ma,
    }

def prisma_flow(prisma: dict):
    n_id = prisma.get("records_identified", 0)
    n_scr = prisma.get("records_screened_ta", 0)
    n_exc = prisma.get("records_excluded_ta", 0)
    n_ft = prisma.get("full_text_assessed", 0)
    n_ft_exc = prisma.get("full_text_excluded", 0)
    n_ma = prisma.get("studies_included_ma", 0)

    st.markdown(
        f"""
<div class="flow">
  <div class="flow-box"><div class="t">Records identified</div><div class="n">PubMed: <b>{n_id}</b></div></div>
  <div class="flow-arrow">↓</div>
  <div class="flow-box"><div class="t">Title/Abstract screened</div><div class="n">Screened: <b>{n_scr}</b> | Excluded: <b>{n_exc}</b></div></div>
  <div class="flow-arrow">↓</div>
  <div class="flow-box"><div class="t">Full-text assessed</div><div class="n">Assessed: <b>{n_ft}</b> | Excluded: <b>{n_ft_exc}</b></div></div>
  <div class="flow-arrow">↓</div>
  <div class="flow-box"><div class="t">Included in meta-analysis</div><div class="n"><b>{n_ma}</b></div></div>
</div>
""",
        unsafe_allow_html=True,
    )
    if prisma.get("full_text_exclusion_reasons"):
        st.caption("Full-text exclusion reasons:")
        st.json(prisma["full_text_exclusion_reasons"])

# =========================================================
# UI: Input + Run
# =========================================================
st.markdown(
    "<div class='notice'><b>工作方式</b><br>"
    "輸入一句話 → 產生 PICO（可展開修正）→ 產生 PubMed 搜尋式（可手改）→ 抓 records → 粗篩（AI 可選/可人工 override）"
    "→ Full text review（含排除理由 + PDF 上傳 + 可選抽字）→ Data extraction 寬表（schema 可改、可一次編輯再 commit）"
    "→ MA + 森林圖（按鈕執行；避免輸入時跳動）→ ROB 2.0 → 稿件草稿（分段顯示 + 可選 BYOK 強化）。"
    "</div>",
    unsafe_allow_html=True,
)

st.markdown("<div class='hr'></div>", unsafe_allow_html=True)

col1, col2 = st.columns([2,1])
with col1:
    st.markdown(f"**{t('question_notice')}**")
    st.text_input(t("question_label"), key="question", help=t("question_help"))
with col2:
    st.markdown("**Run**")
    run_clicked = st.button(t("run"), type="primary")

if st.session_state.get("question") and has_cjk(st.session_state.get("question","")):
    st.info(t("english_hint"))

# Tabs
tabs = st.tabs([
    t("tabs_overview"),
    t("tabs_step1"),
    t("tabs_step2"),
    t("tabs_step34"),
    t("tabs_ft"),
    t("tabs_extract"),
    t("tabs_ma"),
    t("tabs_rob2"),
    t("tabs_ms"),
    t("tabs_diag"),
])

# =========================================================
# Run pipeline (does not force, just populates state)
# =========================================================
def run_pipeline():
    q = st.session_state.get("question","").strip()
    if not q:
        return

    proto, q_en = question_to_protocol_llm(q)
    st.session_state["protocol"] = proto
    st.session_state["question_en"] = q_en

    # Auto build query
    auto_q = build_pubmed_query(proto, st.session_state.get("article_type","不限"), st.session_state.get("custom_pubmed_filter",""))
    st.session_state["pubmed_query_auto"] = auto_q
    # If user hasn't edited manually yet, set pubmed_query to auto
    if not st.session_state.get("pubmed_query"):
        st.session_state["pubmed_query"] = auto_q

    # Fetch pubmed
    df, diag = fetch_pubmed(st.session_state["pubmed_query"], max_records=int(st.session_state.get("max_pubmed_records",1000) or 1000), page_size=int(st.session_state.get("pubmed_page_size",200) or 200))
    st.session_state["pubmed_records"] = df
    st.session_state["diagnostics"] = diag

    # Feasibility: SR/MA/NMA scan (quick)
    feas_q = f"({auto_q}) AND (systematic review[pt] OR meta-analysis[pt] OR \"systematic review\"[tiab] OR \"meta-analysis\"[tiab] OR \"network meta-analysis\"[tiab] OR NMA[tiab])"
    st.session_state["feas_query"] = feas_q
    hits, _ = fetch_pubmed(feas_q, max_records=40, page_size=int(st.session_state.get('pubmed_page_size',200) or 200))
    st.session_state["srma_hits"] = hits

    # Title/abstract AI suggestion
    ta_ai = {}
    ta_reason = {}
    ta_conf = {}
    if not df.empty:
        if llm_available():
            # Batch lightly: first N titles/abstracts to reduce cost
            # (User can rerun later by clearing overrides)
            N = min(50, df.shape[0])
            payload = df.head(N)[["PMID","Title","Abstract","PublicationTypes"]].to_dict(orient="records")
            sys = (
                "You are screening PubMed records for a systematic review. "
                "Given PICO and a list of records, label each record as Include/Exclude/Unsure for FULL-TEXT review, "
                "and provide a 1-2 sentence rationale and a confidence 0-1. "
                "Do not fabricate. Return STRICT JSON: {pmid: {label, reason, confidence}}."
            )
            user = json.dumps({
                "language": st.session_state.get("UI_LANG","zh-TW"),
                "constraints": {
                    "article_type": st.session_state.get("article_type","不限"),
                    "custom_pubmed_filter": st.session_state.get("custom_pubmed_filter",""),
                    "goal_mode": proto.goal_mode,
                    "strict_rct_screen": bool(st.session_state.get("strict_rct_screen", True)),
                },
                "PICO": proto.to_dict().get("pico"),
                "records": payload,
            }, ensure_ascii=False)
            d = llm_json(sys, user, max_tokens=1700) or {}
            for rec in payload:
                pmid = str(rec.get("PMID",""))
                item = d.get(pmid) or d.get(str(pmid)) or {}
                lbl = str(item.get("label") or "Unsure").strip()
                if lbl not in ["Include","Exclude","Unsure"]:
                    lbl = "Unsure"
                ta_ai[pmid] = lbl
                ta_reason[pmid] = str(item.get("reason") or "")
                try:
                    ta_conf[pmid] = float(item.get("confidence") or 0.5)
                except Exception:
                    ta_conf[pmid] = 0.5
            # For remaining beyond N, heuristic
            for _, r in df.iloc[N:].iterrows():
                pmid = str(r.get("PMID",""))
                lbl, rs, cf = heuristic_ta(proto, r)
                ta_ai[pmid] = lbl
                ta_reason[pmid] = rs
                ta_conf[pmid] = cf
        else:
            for _, r in df.iterrows():
                pmid = str(r.get("PMID",""))
                lbl, rs, cf = heuristic_ta(proto, r)
                ta_ai[pmid] = lbl
                ta_reason[pmid] = rs
                ta_conf[pmid] = cf

    st.session_state["ta_ai"] = ta_ai
    st.session_state["ta_ai_reason"] = ta_reason
    st.session_state["ta_ai_conf"] = ta_conf

    # Init full-text decisions for kept records
    for pmid, lbl in ta_ai.items():
        if pmid not in st.session_state["ft_decision"]:
            st.session_state["ft_decision"][pmid] = "Not reviewed"
            st.session_state["ft_reason"][pmid] = ""


def _current_pmids() -> set:
    df = st.session_state.get("pubmed_records")
    if df is None or not isinstance(df, pd.DataFrame) or df.empty or "PMID" not in df.columns:
        return set()
    return set(str(x) for x in df["PMID"].astype(str).tolist() if str(x).strip())


def _prune_dict(d: Dict[str, Any], keep: set) -> Dict[str, Any]:
    if not isinstance(d, dict):
        return {}
    return {k: v for k, v in d.items() if str(k) in keep}


def _prune_after_refetch():
    """Keep user work for records still present; drop the rest to avoid mismatch."""
    keep = _current_pmids()
    st.session_state["ta_ai"] = _prune_dict(st.session_state.get("ta_ai", {}), keep)
    st.session_state["ta_ai_reason"] = _prune_dict(st.session_state.get("ta_ai_reason", {}), keep)
    st.session_state["ta_ai_conf"] = _prune_dict(st.session_state.get("ta_ai_conf", {}), keep)
    st.session_state["ta_override"] = _prune_dict(st.session_state.get("ta_override", {}), keep)
    st.session_state["ta_override_reason"] = _prune_dict(st.session_state.get("ta_override_reason", {}), keep)

    st.session_state["ft_decision"] = _prune_dict(st.session_state.get("ft_decision", {}), keep)
    st.session_state["ft_reason"] = _prune_dict(st.session_state.get("ft_reason", {}), keep)
    st.session_state["ft_pdf"] = _prune_dict(st.session_state.get("ft_pdf", {}), keep)
    st.session_state["ft_text"] = _prune_dict(st.session_state.get("ft_text", {}), keep)

    ex = st.session_state.get("extract_df")
    if isinstance(ex, pd.DataFrame) and not ex.empty and "PMID" in ex.columns:
        st.session_state["extract_df"] = ex[ex["PMID"].astype(str).isin(keep)].reset_index(drop=True)

    rb = st.session_state.get("rob2_df")
    if isinstance(rb, pd.DataFrame) and not rb.empty and "PMID" in rb.columns:
        st.session_state["rob2_df"] = rb[rb["PMID"].astype(str).isin(keep)].reset_index(drop=True)


def _llm_enabled() -> bool:
    return bool(
        st.session_state.get("byok_enabled")
        and st.session_state.get("byok_consent")
        and (st.session_state.get("byok_base_url") or "").strip()
        and (st.session_state.get("byok_model") or "").strip()
        and (st.session_state.get("byok_key") or "").strip()
    )


def compute_ta_suggestions(proto: Protocol, df: pd.DataFrame):
    """Compute AI (or heuristic) title/abstract screening suggestions for FULL-TEXT."""
    ta_ai: Dict[str, str] = {}
    ta_reason: Dict[str, str] = {}
    ta_conf: Dict[str, float] = {}

    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        st.session_state["ta_ai"] = {}
        st.session_state["ta_ai_reason"] = {}
        st.session_state["ta_ai_conf"] = {}
        return

    # Build payload for LLM
    payload = []
    for _, r in df.iterrows():
        payload.append({
            "PMID": str(r.get("PMID","")),
            "Title": str(r.get("Title","") or ""),
            "Abstract": str(r.get("Abstract","") or "")[:3500],
            "Year": str(r.get("Year","") or ""),
            "FirstAuthor": str(r.get("First_author","") or r.get("FirstAuthor","") or ""),
            "Journal": str(r.get("Journal","") or ""),
            "PublicationTypes": str(r.get("PublicationTypes","") or ""),
        })

    if _llm_enabled():
        try:
            sys = (
                "You are performing title/abstract screening for an evidence synthesis. "
                "Given (a) PICO, (b) constraints, and (c) a list of PubMed records, you must label EACH record as "
                "Include / Exclude / Unsure for FULL-TEXT review. "
                "Be conservative (prefer Unsure over wrong Exclude) unless a constraint is clearly violated. "
                "For the rationale, write 1–3 short bullet points that are audit-friendly: "
                "(1) which PICO elements match (mention exact terms found), "
                "(2) study design check (e.g., RCT yes/no and why; refer to PublicationTypes/title/abstract cues), "
                "(3) any hard exclusion reason if applicable. "
                "Return STRICT JSON only with schema: {pmid: {label: str, confidence: float, reason: str}}. "
                "Do not invent data not present in the record."
            )
            user = json.dumps({"PICO": proto.to_dict().get("pico"), "records": payload}, ensure_ascii=False)
            d = llm_json(sys, user, max_tokens=1700) or {}
            for rec in payload:
                pmid = str(rec.get("PMID",""))
                item = d.get(pmid) or d.get(str(pmid)) or {}
                lbl = str(item.get("label") or "Unsure").strip()
                if lbl not in ["Include","Exclude","Unsure"]:
                    lbl = "Unsure"
                ta_ai[pmid] = lbl
                ta_reason[pmid] = str(item.get("reason") or "").strip()
                try:
                    cf = float(item.get("confidence", 0.5))
                except Exception:
                    cf = 0.5
                ta_conf[pmid] = max(0.0, min(1.0, cf))
        except Exception as e:
            # Fall back to heuristic
            for _, r in df.iterrows():
                pmid = str(r.get("PMID",""))
                lbl, rs, cf = heuristic_ta(proto, r)
                ta_ai[pmid] = lbl
                ta_reason[pmid] = rs
                ta_conf[pmid] = cf
    else:
        for _, r in df.iterrows():
            pmid = str(r.get("PMID",""))
            lbl, rs, cf = heuristic_ta(proto, r)
            ta_ai[pmid] = lbl
            ta_reason[pmid] = rs
            ta_conf[pmid] = cf

    st.session_state["ta_ai"] = ta_ai
    st.session_state["ta_ai_reason"] = ta_reason
    st.session_state["ta_ai_conf"] = ta_conf


def refetch_pubmed_and_sync():
    """Re-fetch PubMed using the (possibly manually edited) query and re-sync downstream steps."""
    q = (st.session_state.get("pubmed_query") or "").strip()
    if not q:
        return
    max_n = int(st.session_state.get("max_pubmed_records", 1000) or 1000)
    page_size = int(st.session_state.get("pubmed_page_size", 200) or 200)

    df, diag = fetch_pubmed(q, max_records=max_n, page_size=page_size)
    st.session_state["pubmed_records"] = df
    st.session_state["diagnostics"] = diag

    # Feasibility query should follow CURRENT query (not the old auto_q)
    feas_q = f"({q}) AND (systematic review[pt] OR meta-analysis[pt] OR \"systematic review\"[tiab] OR \"meta analysis\"[tiab] OR \"network meta-analysis\"[tiab] OR NMA[tiab])"
    st.session_state["feas_query"] = feas_q
    hits, _ = fetch_pubmed(feas_q, max_records=60, page_size=page_size)
    st.session_state["srma_hits"] = hits

    proto = st.session_state.get("protocol")
    if proto and isinstance(proto, Protocol):
        compute_ta_suggestions(proto, df)

    # Keep only work that still matches current PMIDs
    _prune_after_refetch()

if run_clicked:
    run_pipeline()

# =========================================================
# Tab 0: Overview / PRISMA
# =========================================================
with tabs[0]:
    st.subheader(t("tabs_overview"))
    prisma = compute_prisma(st.session_state.get("pubmed_records"))
    prisma_flow(prisma)

# =========================================================
# Tab 1: Step 1 - Query editable + PICO edits
# =========================================================
with tabs[1]:
    st.subheader(t("tabs_step1"))

    proto: Protocol = st.session_state.get("protocol")
    if not proto:
        st.info("請先輸入問題並按 Run。")
    else:
        with st.expander(t("pico_edit"), expanded=True):
            # Manual PICO correction (core ask)
            c1, c2 = st.columns(2)
            with c1:
                P = st.text_input("P", value=proto.P, key="edit_P")
                I_ = st.text_input("I", value=proto.I, key="edit_I")
                C = st.text_input("C", value=proto.C, key="edit_C")
                O = st.text_input("O", value=proto.O, key="edit_O")
                NOT = st.text_input("NOT", value=proto.NOT, key="edit_NOT")
            with c2:
                st.caption("Synonyms (comma-separated; English preferred)")
                P_syn = st.text_area("P synonyms", value=", ".join(proto.P_syn or []), key="edit_P_syn", height=70)
                I_syn = st.text_area("I synonyms", value=", ".join(proto.I_syn or []), key="edit_I_syn", height=70)
                C_syn = st.text_area("C synonyms", value=", ".join(proto.C_syn or []), key="edit_C_syn", height=70)
                O_syn = st.text_area("O synonyms", value=", ".join(proto.O_syn or []), key="edit_O_syn", height=70)
                st.caption("MeSH candidates (comma-separated)")
                mesh_P = st.text_area("MeSH P", value=", ".join(proto.mesh_P or []), key="edit_mesh_P", height=60)
                mesh_I = st.text_area("MeSH I", value=", ".join(proto.mesh_I or []), key="edit_mesh_I", height=60)
                mesh_C = st.text_area("MeSH C", value=", ".join(proto.mesh_C or []), key="edit_mesh_C", height=60)
                mesh_O = st.text_area("MeSH O", value=", ".join(proto.mesh_O or []), key="edit_mesh_O", height=60)

            if st.button(t("pico_apply"), type="primary"):
                proto.P = norm_text(P); proto.I = norm_text(I_); proto.C = norm_text(C); proto.O = norm_text(O)
                proto.NOT = norm_text(NOT)
                proto.P_syn = [norm_text(x) for x in re.split(r"[,\n]+", P_syn or "") if norm_text(x)]
                proto.I_syn = [norm_text(x) for x in re.split(r"[,\n]+", I_syn or "") if norm_text(x)]
                proto.C_syn = [norm_text(x) for x in re.split(r"[,\n]+", C_syn or "") if norm_text(x)]
                proto.O_syn = [norm_text(x) for x in re.split(r"[,\n]+", O_syn or "") if norm_text(x)]
                proto.mesh_P = [norm_text(x) for x in re.split(r"[,\n]+", mesh_P or "") if norm_text(x)]
                proto.mesh_I = [norm_text(x) for x in re.split(r"[,\n]+", mesh_I or "") if norm_text(x)]
                proto.mesh_C = [norm_text(x) for x in re.split(r"[,\n]+", mesh_C or "") if norm_text(x)]
                proto.mesh_O = [norm_text(x) for x in re.split(r"[,\n]+", mesh_O or "") if norm_text(x)]
                st.session_state["protocol"] = proto

                # rebuild query auto + sync current query to auto (but keep user's manual edits if present)
                auto_q = build_pubmed_query(proto, st.session_state.get("article_type","不限"), st.session_state.get("custom_pubmed_filter",""))
                st.session_state["pubmed_query_auto"] = auto_q
                if st.session_state.get("pubmed_query","") == "" or st.session_state.get("pubmed_query","") == st.session_state.get("pubmed_query_auto",""):
                    st.session_state["pubmed_query"] = auto_q
                st.success("已套用 PICO 修正；請到下方檢查/手改 PubMed query。")

        st.markdown("**Protocol (current)**")
        st.code(pretty_json(proto.to_dict()), language="json")

        st.markdown("### " + t("pubmed_edit"))
        st.text_area("", key="pubmed_query", height=120)
        cA, cB, cC = st.columns([1,1,2])
        with cA:
            if st.button(t("pubmed_refetch"), type="primary"):
                refetch_pubmed_and_sync()
                df = st.session_state.get("pubmed_records", pd.DataFrame())
                diag = st.session_state.get("diagnostics", {})
                st.success(f"抓到 {df.shape[0]} 篇（PubMed count={diag.get('pubmed_total_count',0)}）。")
                st.info("已使用手動 PubMed query 重新抓取並同步後續步驟；若先前已有人工標記/抽取，系統僅保留仍存在於目前 records 的部分。")
        with cB:
            if st.button(t("pubmed_restore")):
                st.session_state["pubmed_query"] = st.session_state.get("pubmed_query_auto","")
        with cC:
            st.download_button(
                t("download_query"),
                data=st.session_state.get("pubmed_query","").encode("utf-8"),
                file_name="pubmed_query.txt",
                mime="text/plain",
            )

# =========================================================
# Tab 2: Feasibility scan + recommendations
# =========================================================
    st.markdown("---")
    st.markdown("### PRISMA（即時總覽）")
    prisma_now = compute_prisma(st.session_state.get("pubmed_records"))
    prisma_flow(prisma_now)
    st.caption("此 PRISMA 會隨 Step 3/4 初篩、Step 4b 全文決策、Step 5/6 的更新而自動變化。若你在 Step 1 手改搜尋式，請記得按 Refetch 重新抓 records。")

with tabs[2]:
    st.subheader(t("feas_title"))
    hits = st.session_state.get("srma_hits")
    if hits is None or not isinstance(hits, pd.DataFrame) or hits.empty:
        st.info("尚未執行可行性掃描。請先 Run。")
    else:
        st.caption("此區塊用來快速判斷：是否已有大量 SR/MA/NMA；是否需要裁切題目（族群/介入/比較/結果/研究設計）。")
        st.markdown("**Feasibility query (auto)**")
        st.code(st.session_state.get("feas_query",""), language="text")
        st.markdown("**Existing SR/MA/NMA hits (sample)**")
        show_cols = ["PMID","Year","First_author","Journal","Title"]
        hits2 = ensure_columns(hits.copy(), show_cols, default="")
        st.dataframe(hits2[show_cols].head(20), use_container_width=True)

        # Heuristic recommendations
        n_hits = hits.shape[0]
        recs = []
        if n_hits >= 15:
            recs.append("已有相當多 SR/MA/NMA：建議縮小題目（特定族群/特定 IOL 型號/特定 outcome/特定追蹤時間/只納入 RCT）。")
            recs.append("若仍要做：可考慮『更新版』（加上近 2-3 年新 RCT）、或做 subgroup/診斷/手術方式差異。")
        elif 5 <= n_hits < 15:
            recs.append("已有一些 SR/MA：建議先讀最相關 2-3 篇，確認缺口（outcome 未涵蓋、亞組未做、研究設計/新器材）。")
        else:
            recs.append("目前 SR/MA/NMA 命中不多：可能具有可行性。仍建議檢查是否需要加入 MeSH/型號關鍵字提高召回。")

        if st.session_state.get("goal_mode") == "Fast / feasible (gap-fill)":
            recs.append("目標取向為『快速可行』：可優先選擇最容易得到 RCT、outcome 定義一致的題目。")
        else:
            recs.append("目標取向為『嚴謹範圍』：建議預先定義 outcome/timepoint 並限制納入條件，避免不可比。")

        st.markdown("**綜合建議**")
        for r in recs:
            st.write("- " + r)

        if llm_available():
            with st.expander(t("feas_optional"), expanded=False):
                if st.button("Generate feasibility report (BYOK)"):
                    proto: Protocol = st.session_state.get("protocol")
                    sys = (
                        "You are helping plan a feasible SR/MA/NMA. "
                        "Given the research question, PICO, and a sample of existing SR/MA/NMA hits, "
                        "produce a feasibility report with (1) summary, (2) whether to proceed, "
                        "(3) recommended PICO refinements (narrow/shift), and (4) suggested outcomes/timepoints. "
                        "Return STRICT JSON with keys: proceed (yes/no), summary, recommended_changes (list), suggested_outcomes (list)."
                    )
                    sample = hits2[show_cols].head(12).to_dict(orient="records")
                    user = json.dumps({"question_en": st.session_state.get("question_en",""),
                                      "PICO": proto.to_dict().get("pico"),
                                      "existing_srma_sample": sample}, ensure_ascii=False)
                    d = llm_json(sys, user, max_tokens=1200) or {}
                    st.json(d)

# =========================================================
# Tab 3: Step 3+4 - Records + screening (merged)
# =========================================================

with tabs[3]:
    st.subheader(t("tabs_step34"))
    df = st.session_state.get("pubmed_records")

    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        st.info(t("records_none"))
    else:
        # Ensure stable columns (avoid KeyError when upstream schema changes)
        df = ensure_columns(
            df.copy(),
            ["PMID", "Title", "Abstract", "Year", "First_author", "Journal", "DOI", "PMCID", "PublicationTypes"],
            default="",
        )

        st.caption(
            "此步驟是『初篩（Title/Abstract）』：AI 只是在旁邊提供建議與理由，最終納入/排除以你在此步驟的決策為準。"
            "若已在 Settings 選擇文章類型（如 RCT），AI/規則會把『不符合設計』的研究優先標示為 Exclude。"
        )

        st.session_state.setdefault("view_mode_step34", "卡片")
        st.session_state.setdefault("strict_rct_screen", True)

        # Bulk / utilities
        c1, c2, c3, c4, c5 = st.columns([1.35, 1.2, 1.0, 1.15, 2.3])
        with c1:
            if st.button("重新計算 AI 初篩建議（依最新 PICO/限制）"):
                proto = st.session_state.get("protocol")
                if proto and isinstance(proto, Protocol):
                    compute_ta_suggestions(proto, df)
                    st.success("已更新 AI 建議。")
        with c2:
            if st.button("Override all Unsure → Include"):
                for pmid in df["PMID"].astype(str).tolist():
                    if st.session_state.get("ta_ai", {}).get(pmid, "Unsure") == "Unsure":
                        st.session_state["ta_override"][pmid] = "Include"
                        st.session_state["ta_override_reason"][pmid] = "Batch override: keep for full text."
        with c3:
            if st.button("Clear all overrides"):
                st.session_state["ta_override"] = {}
                st.session_state["ta_override_reason"] = {}
        with c4:
            st.session_state["strict_rct_screen"] = st.checkbox(
                "嚴格依 RCT 篩（已選 RCT 時）",
                value=bool(st.session_state.get("strict_rct_screen", True)),
                help="開啟時：若已選 RCT，且摘要/PublicationTypes 無 RCT 線索，會更傾向 Exclude。",
            )
        with c5:
            view_mode = st.radio(
                "顯示方式",
                options=["卡片", "表格（精簡，可改 Include/Exclude）"],
                horizontal=True,
                index=0 if st.session_state.get("view_mode_step34", "卡片") == "卡片" else 1,
                key="view_mode_step34",
            )

        def _final_label(pmid: str) -> str:
            ov = (st.session_state.get("ta_override", {}) or {}).get(pmid, "")
            if ov:
                return ov
            return (st.session_state.get("ta_ai", {}) or {}).get(pmid, "Unsure") or "Unsure"

        # Table view (editable)
        show_cards = (st.session_state.get("view_mode_step34", "卡片") == "卡片")
        if not show_cards:
            view = df[["PMID", "Year", "First_author", "Journal", "PublicationTypes", "Title"]].copy()
            view["AI_suggest"] = [st.session_state.get("ta_ai", {}).get(str(p), "Unsure") for p in view["PMID"].astype(str)]
            view["AI_conf"] = [float(st.session_state.get("ta_ai_conf", {}).get(str(p), 0.5) or 0.5) for p in view["PMID"].astype(str)]
            view["AI_reason"] = [st.session_state.get("ta_ai_reason", {}).get(str(p), "") for p in view["PMID"].astype(str)]
            view["Final"] = [_final_label(str(p)) for p in view["PMID"].astype(str)]
            view["Override_reason"] = [st.session_state.get("ta_override_reason", {}).get(str(p), "") for p in view["PMID"].astype(str)]

            edited = st.data_editor(
                view,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Final": st.column_config.SelectboxColumn(
                        "Final (你決定)",
                        options=["Include", "Unsure", "Exclude"],
                        required=True,
                    ),
                    "AI_conf": st.column_config.NumberColumn("AI_conf", format="%.2f"),
                    "AI_reason": st.column_config.TextColumn("AI_reason", width="large"),
                    "Override_reason": st.column_config.TextColumn("Override_reason（可留空）", width="large"),
                },
                disabled=["AI_suggest", "AI_conf", "AI_reason"],
            )

            cA, cB = st.columns([1, 3])
            with cA:
                if st.button("套用表格修改", type="primary"):
                    ta_ai = st.session_state.get("ta_ai", {}) or {}
                    for _, r in edited.iterrows():
                        pmid = str(r.get("PMID",""))
                        final = str(r.get("Final","Unsure")).strip() or "Unsure"
                        or_reason = str(r.get("Override_reason","") or "").strip()
                        ai0 = ta_ai.get(pmid, "Unsure") or "Unsure"

                        if final == ai0 and not or_reason:
                            # identical to AI, drop override to keep state clean
                            st.session_state["ta_override"].pop(pmid, None)
                            st.session_state["ta_override_reason"].pop(pmid, None)
                        else:
                            st.session_state["ta_override"][pmid] = final
                            if or_reason:
                                st.session_state["ta_override_reason"][pmid] = or_reason
                            else:
                                st.session_state["ta_override_reason"][pmid] = f"Override: set to {final} in table."
                    st.success("已更新初篩決策（Final）。PRISMA/全文階段會自動跟著變化。")
            with cB:
                # Summary counts
                counts = {"Include": 0, "Unsure": 0, "Exclude": 0}
                for pmid in df["PMID"].astype(str).tolist():
                    counts[_final_label(pmid)] = counts.get(_final_label(pmid), 0) + 1
                st.caption(f"目前 Final：Include {counts.get('Include',0)}｜Unsure {counts.get('Unsure',0)}｜Exclude {counts.get('Exclude',0)}")

        # Card view (grouped)
        else:
            def badge_html(label: str) -> str:
                label = (label or "").strip()
                styles = {
                    "Include": ("#0f5132", "#d1e7dd"),
                    "Unsure": ("#664d03", "#fff3cd"),
                    "Exclude": ("#842029", "#f8d7da"),
                }
                fg, bg = styles.get(label, ("#1f2328", "#e9ecef"))
                return f"<span style='display:inline-block;padding:2px 8px;border-radius:999px;font-size:.85rem;font-weight:600;color:{fg};background:{bg};border:1px solid rgba(0,0,0,0.08)'> {label} </span>"

            # Build groups by FINAL label (override takes precedence)
            groups: Dict[str, List[dict]] = {"Include": [], "Unsure": [], "Exclude": []}
            for _, r in df.iterrows():
                pmid = str(r.get("PMID","") or "")
                lbl = _final_label(pmid)
                if lbl not in groups:
                    lbl = "Unsure"
                groups[lbl].append(r.to_dict())

            def render_record_card(r):
                pmid = str(r.get("PMID", "") or "").strip()
                title = str(r.get("Title", "") or "").strip()
                abstract = str(r.get("Abstract", "") or "")
                year = str(r.get("Year", "") or "").strip()
                fa = str(r.get("First_author", "") or "").strip()
                journal = str(r.get("Journal", "") or "").strip()
                doi = str(r.get("DOI", "") or "").strip()
                pmcid = str(r.get("PMCID", "") or "").strip()
                pubtypes = str(r.get("PublicationTypes", "") or "").strip()

                ai_lbl = (st.session_state.get("ta_ai", {}) or {}).get(pmid, "Unsure") or "Unsure"
                ai_reason = (st.session_state.get("ta_ai_reason", {}) or {}).get(pmid, "")
                try:
                    ai_conf = float((st.session_state.get("ta_ai_conf", {}) or {}).get(pmid, 0.5) or 0.5)
                except Exception:
                    ai_conf = 0.5

                ov = (st.session_state.get("ta_override", {}) or {}).get(pmid, "")
                ov_reason = (st.session_state.get("ta_override_reason", {}) or {}).get(pmid, "")
                final_lbl = ov if ov else ai_lbl

                head = f"PMID:{pmid or '—'}｜{short(title or '—', 110)}"
                with st.expander(head, expanded=False):
                    st.markdown("<div class='card'>", unsafe_allow_html=True)

                    meta = f"PMID: {pmid or '—'}　|　DOI: {doi or '—'}　|　Year: {year or '—'}　|　First author: {fa or '—'}　|　Journal: {journal or '—'}"
                    st.markdown(f"**{title or '—'}**")
                    st.markdown(f"<div class='small'>{meta}</div>", unsafe_allow_html=True)
                    if pubtypes:
                        st.markdown(f"<div class='small'>PublicationTypes: {html.escape(pubtypes)}</div>", unsafe_allow_html=True)

                    links = []
                    if pubmed_link(pmid):
                        links.append(f"[PubMed]({maybe_ezproxy(pubmed_link(pmid))})")
                    if doi:
                        links.append(f"[DOI]({maybe_ezproxy(doi_link(doi))})")
                    if pmcid:
                        links.append(f"[PMC]({maybe_ezproxy(pmc_link(pmcid))})")
                    if resolver_url(doi, pmid):
                        links.append(f"[學院全文連結]({resolver_url(doi, pmid)})")
                    if links:
                        st.markdown(" | ".join(links))

                    st.markdown("**Abstract**")
                    st.write(abstract if abstract else "(no abstract)")

                    st.markdown("**AI 建議**")
                    st.markdown(f"{badge_html(ai_lbl)}　信心度：{ai_conf:.2f}", unsafe_allow_html=True)
                    st.write(ai_reason or "(no AI reason)")

                    st.markdown("**Final（你決定）**")
                    c1, c2 = st.columns([1, 2.2])
                    with c1:
                        picked = st.selectbox(
                            "Final label",
                            options=["", "Include", "Unsure", "Exclude"],
                            index=0,
                            key=f"ta_final_{pmid}",
                            help="選了才會寫入 override；空白=沿用 AI。",
                        )
                    with c2:
                        reason0 = st.text_input("Override reason (optional)", value=ov_reason, key=f"ta_why_{pmid}")

                    if st.button("套用此篇 Final", key=f"ta_apply_{pmid}", type="primary"):
                        if picked:
                            st.session_state["ta_override"][pmid] = picked
                            st.session_state["ta_override_reason"][pmid] = reason0 or f"Override: {picked}"
                        else:
                            st.session_state["ta_override"].pop(pmid, None)
                            st.session_state["ta_override_reason"].pop(pmid, None)
                        st.success("已更新此篇 Final。")

                    st.markdown("</div>", unsafe_allow_html=True)

            # Render by groups
            for lbl in ["Include", "Unsure", "Exclude"]:
                items = groups.get(lbl, [])
                st.markdown(f"### {badge_html(lbl)} {lbl}（{len(items)}）", unsafe_allow_html=True)
                if not items:
                    st.caption("(none)")
                for r in items:
                    render_record_card(r)

# Tab 4: Step 4b Full text review (decisions + reasons + PDF upload)
# =========================================================
# Tab 4: Step 4b Full text review (decisions + reasons + PDF upload)
# =========================================================
FULLTEXT_EXCLUSION_REASONS = [
    "Not relevant population",
    "Not relevant intervention/comparator",
    "Not comparative (single arm)",
    "Wrong study design",
    "Duplicate/overlap",
    "No usable outcome data",
    "Conference abstract only",
    "Full text not accessible",
    "Other",
]


with tabs[4]:
    st.subheader(t("tabs_ft"))
    df = st.session_state.get("pubmed_records")
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        st.info("請先 Run 並抓到 records。")
    else:
        df = ensure_columns(df.copy(), ["PMID","Title","Year","First_author","Journal","DOI","PublicationTypes"], default="")

        # Show only records kept after TA (Include/Unsure)
        def final_ta(pmid: str) -> str:
            return st.session_state.get("ta_override",{}).get(pmid) or st.session_state.get("ta_ai",{}).get(pmid,"Unsure")

        kept = df[df["PMID"].astype(str).apply(lambda x: final_ta(str(x)) != "Exclude")].copy()
        st.caption(f"進入全文階段的 records（初篩未排除）：{kept.shape[0]} 篇")
        if kept.empty:
            st.info("沒有可做全文審查的 record。")
        else:
            st.session_state.setdefault("ft_decision", {})
            st.session_state.setdefault("ft_reason", {})
            st.session_state.setdefault("ft_pdf", {})
            st.session_state.setdefault("ft_text", {})

            # Bulk upload PDFs (optional)
            with st.expander(t("ft_bulk_upload"), expanded=False):
                st.caption("注意：若為校內訂閱/付費期刊全文，請勿上傳到雲端部署（授權風險）。建議只上傳 OA/PMC 或本機版使用。")
                uploads = st.file_uploader("Upload PDFs (multiple)", type=["pdf"], accept_multiple_files=True)
                if uploads:
                    for up in uploads:
                        name = up.name
                        m = re.search(r"(\d{6,9})", name)
                        if not m:
                            continue
                        pmid = m.group(1)
                        st.session_state["ft_pdf"][pmid] = up.getvalue()
                    st.success("已儲存上傳的 PDF（以檔名中的 PMID 對應）。")

            # Select one study to work on (prevents state loss / too many widgets)
            kept["Label"] = kept.apply(lambda r: f"{r['PMID']}｜{short(str(r['Title'] or ''), 90)}", axis=1)
            options = kept["Label"].tolist()
            default_idx = 0
            current = st.selectbox("目前處理的全文研究", options=options, index=default_idx, key="ft_current_pick")
            pmid = current.split("｜")[0].strip()

            rr = kept[kept["PMID"].astype(str) == pmid]
            r0 = rr.iloc[0].to_dict() if not rr.empty else {}
            title = str(r0.get("Title","") or "")
            year = str(r0.get("Year","") or "")
            fa = str(r0.get("First_author","") or "")
            journal = str(r0.get("Journal","") or "")
            doi = str(r0.get("DOI","") or "")
            pubtypes = str(r0.get("PublicationTypes","") or "")

            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown(f"**{title or '—'}**")
            st.markdown(f"<span class='small'>PMID: {pmid} | Year: {year or '—'} | First author: {fa or '—'} | Journal: {journal or '—'}</span>", unsafe_allow_html=True)
            if pubtypes:
                st.markdown(f"<div class='small'>PublicationTypes: {html.escape(pubtypes)}</div>", unsafe_allow_html=True)

            links = []
            if pubmed_link(pmid):
                links.append(f"[PubMed]({maybe_ezproxy(pubmed_link(pmid))})")
            if doi:
                links.append(f"[DOI]({maybe_ezproxy(doi_link(doi))})")
            if resolver_url(doi, pmid):
                links.append(f"[學院全文連結]({resolver_url(doi, pmid)})")
            if links:
                st.markdown(" | ".join(links))

            # Full-text decision
            cA, cB, cC = st.columns([1.1, 1.5, 1.4])
            with cA:
                decision_new = st.selectbox(
                    "Full-text decision",
                    options=["Not reviewed","Include for meta-analysis","Exclude"],
                    index=["Not reviewed","Include for meta-analysis","Exclude"].index(st.session_state["ft_decision"].get(pmid, "Not reviewed")),
                    key=f"ft_dec_{pmid}",
                )
            with cB:
                reason_new = st.selectbox(
                    "If Exclude: reason",
                    options=[""] + FULLTEXT_EXCLUSION_REASONS,
                    index=0,
                    key=f"ft_reason_sel_{pmid}",
                )
            with cC:
                other_reason = st.text_input("Other details (optional)", key=f"ft_reason_free_{pmid}")

            if st.button("Save decision", type="primary", key=f"ft_save_{pmid}"):
                st.session_state["ft_decision"][pmid] = decision_new
                if decision_new == "Exclude":
                    rr0 = reason_new or "Other"
                    if rr0 == "Other" and other_reason.strip():
                        rr0 = f"Other: {other_reason.strip()}"
                    st.session_state["ft_reason"][pmid] = rr0
                else:
                    st.session_state["ft_reason"][pmid] = ""
                st.success("已儲存全文決策。")

            # PDF upload for this PMID
            up_one = st.file_uploader("Upload PDF for this PMID (optional)", type=["pdf"], key=f"ft_pdf_one_{pmid}")
            if up_one:
                st.session_state["ft_pdf"][pmid] = up_one.getvalue()
                st.success("已儲存此篇 PDF。")

            # Extract / paste full-text
            c1, c2 = st.columns([1, 1])
            with c1:
                if st.button("從 PDF 抽取文字（不含 OCR）", key=f"ft_extract_{pmid}"):
                    pdfb = st.session_state.get("ft_pdf", {}).get(pmid)
                    if not pdfb:
                        st.warning("尚未有此篇 PDF。")
                    else:
                        txt = extract_pdf_text(pdfb, max_pages=30)
                        if not txt.strip():
                            st.warning("抽字結果為空。可能是掃描 PDF（需要 OCR）或版面不支援。建議先 OCR 再貼/上傳文字。")
                        else:
                            # only overwrite if non-empty, to avoid accidental loss
                            st.session_state["ft_text"][pmid] = txt
                            st.success(f"已抽取文字（長度 {len(txt)}）。")
            with c2:
                if st.button("清空此篇全文文字", key=f"ft_clear_{pmid}"):
                    st.session_state["ft_text"][pmid] = ""

            ft_text = st.session_state.get("ft_text", {}).get(pmid, "")
            ft_text = st.text_area(t("ft_text_area"), value=ft_text, height=220, key=f"ft_text_area_{pmid}")
            st.session_state["ft_text"][pmid] = ft_text

            if llm_available():
                with st.expander("AI 閱讀回填（Step 5/ROB 2.0 會用到；請人工核對）", expanded=False):
                    if st.button(t("ft_ai_fill"), key=f"ft_ai_btn_{pmid}", type="primary"):
                        proto: Protocol = st.session_state.get("protocol")
                        schema_cols = [c.strip() for c in (st.session_state.get("extract_schema_text") or default_extraction_schema()).splitlines() if c.strip()]
                        prompt = build_data_extraction_prompt(proto, schema_cols, full_text=ft_text)
                        sys = (
                            "You are an evidence extraction assistant. "
                            "Extract ONLY what is explicitly supported by the provided text. "
                            "If a field is not reported, leave it empty. Return JSON only."
                        )
                        d = llm_json(sys, prompt, max_tokens=1700) or {}
                        # store extraction into wide table (append)
                        if isinstance(d, dict) and schema_cols:
                            row = {c: "" for c in schema_cols}
                            for c in schema_cols:
                                if c in d:
                                    row[c] = d.get(c, "")
                            # Force IDs
                            row["PMID"] = pmid
                            row["First_author"] = fa
                            row["Year"] = year
                            row["Title"] = title
                            row["DOI"] = doi
                            # Append
                            ex = st.session_state.get("extract_df")
                            if not isinstance(ex, pd.DataFrame) or ex.empty:
                                st.session_state["extract_df"] = pd.DataFrame([row])
                            else:
                                st.session_state["extract_df"] = pd.concat([ex, pd.DataFrame([row])], ignore_index=True)
                            st.success("已把 AI 抽取結果追加到 Step 5 寬表（可再人工修正）。")
                        else:
                            st.warning("AI 回填未產生可用 JSON。")

            st.markdown("</div>", unsafe_allow_html=True)

            # Optional: quick list of all kept studies (no full-text widgets)
            with st.expander("查看目前全文階段清單（不含全文輸入框）", expanded=False):
                view = kept[["PMID","Year","First_author","Journal","Title"]].copy()
                view["TA_Final"] = [final_ta(str(p)) for p in view["PMID"].astype(str)]
                view["FT_decision"] = [st.session_state.get("ft_decision",{}).get(str(p),"Not reviewed") for p in view["PMID"].astype(str)]
                st.dataframe(view, use_container_width=True, hide_index=True)

# =========================================================
# Tab 5: Extraction wide table + editor
# =========================================================

with tabs[5]:

    st.subheader(t("tabs_extract"))

    df = st.session_state.get("pubmed_records")
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        st.info("請先 Run 並抓到 records。")
    else:
        # Schema (editable)
        if not st.session_state.get("extract_schema_text"):
            st.session_state["extract_schema_text"] = default_extraction_schema()

        st.markdown("#### " + t("extract_schema"))
        st.caption("目標要求：extraction table 不要寫死 → 在 PICO 層級自主規劃欄位。你可自行增刪欄位，並考量：既有 SR/MA/NMA、既有 RCT 的 primary/secondary outcomes。")
        schema_text = st.text_area("", key="extract_schema_text", height=170)

        schema_cols = [c.strip() for c in (schema_text or "").splitlines() if c.strip()]
        if not schema_cols:
            schema_cols = [c.strip() for c in default_extraction_schema().splitlines() if c.strip()]
            st.warning("schema 為空，已回復預設欄位。")

        st.session_state["extract_df"] = build_extraction_df_from_schema(schema_cols)

        if llm_available() and st.button("BYOK：根據 PICO + 既有 SR/MA 建議 extraction schema"):
            proto: Protocol = st.session_state.get("protocol")
            sys = (
                "You design a data extraction sheet for a systematic review/meta-analysis. "
                "Given PICO and goal mode, propose a list of extraction columns at PICO level. "
                "Must consider: (a) whether prior SR/MA/NMA exists; (b) RCT primary/secondary outcomes to capture; "
                "and (c) effect size and CI fields for meta-analysis. "
                "Return STRICT JSON with key 'columns' as an array of strings."
            )
            user = json.dumps({
                "PICO": proto.to_dict().get("pico"),
                "goal_mode": proto.goal_mode,
                "existing_srma_n": int((st.session_state.get("srma_hits") or pd.DataFrame()).shape[0]),
            }, ensure_ascii=False)
            d = llm_json(sys, user, max_tokens=700) or {}
            cols = d.get("columns") or []
            if isinstance(cols, list) and cols:
                st.session_state["extract_schema_text"] = "\n".join([str(x).strip() for x in cols if str(x).strip()])
                st.success("已更新 schema（你仍可手動調整）。")

        st.markdown("---")
        st.markdown("#### " + t("extract_quick_add"))
        st.caption("用表單一次輸入完再寫入，避免每打一格就 rerun 造成『跳掉』。")

        # Eligible studies: those included for MA (full-text)
        kept_df = ensure_columns(df.copy(), ["PMID","Title","First_author","Year"], default="")
        include_pmids = [pmid for pmid, dec in (st.session_state.get("ft_decision") or {}).items() if dec == "Include for meta-analysis"]

        # 僅顯示 Full-text = Include for meta-analysis（符合 MA 抽取流程）
        pool_df = kept_df.copy()
        if include_pmids:
            pool_df = kept_df[kept_df["PMID"].astype(str).isin(set(map(str, include_pmids)))].copy()
        else:
            st.warning("目前沒有 Full-text decision = Include for meta-analysis 的研究；請先到 Step 4b 完成全文決策。")
            # 進階：允許先用粗篩納入建立抽取表（預設關閉，避免與規範不一致）
            allow_pre_ft = st.checkbox("（進階）尚未完成 Full text 時，暫時顯示 Title/Abstract 納入以先建立抽取表", value=False, key="step5_allow_pre_ft")
            if allow_pre_ft:
                pool_df = kept_df.copy()
            else:
                pool_df = kept_df.iloc[0:0].copy()

        options = []
        pmid_to_row = {}
        for _, r in pool_df.iterrows():
            pmid = str(r["PMID"])
            label = f"{pmid} | {r['First_author']} | {short(r['Title'], 60)}"
            pmid_to_row[label] = r.to_dict()
            options.append(label)

        with st.form("quick_add_form"):
            pick = st.selectbox("Choose record", options=options if options else ["（無可用研究：請先在 Step 4b 設為 Include for meta-analysis）"], index=0)
            picked = pmid_to_row.get(pick, {})
            st.write(f"Selected: PMID={picked.get('PMID','')} | First author={picked.get('First_author','')} | Year={picked.get('Year','')}")
            # minimal fields
            out_label = st.text_input("OutcomeLabel", value=st.session_state.get("ma_outcome_input",""))
            eff_measure = st.selectbox("Effect_measure", options=["OR","RR","HR","MD","SMD"], index=0)
            eff = st.text_input("Effect", value="")
            lcl = st.text_input("Lower_CI", value="")
            ucl = st.text_input("Upper_CI", value="")
            notes = st.text_input("Notes", value="")
            submitted = st.form_submit_button("Append row")

        if submitted and options:
            row = {c: "" for c in schema_cols}
            # Auto-fill bibliographic identifiers so you know which paper you're editing
            row["PMID"] = str(picked.get("PMID",""))
            row["First_author"] = str(picked.get("First_author",""))
            row["Year"] = str(picked.get("Year",""))
            row["Title"] = str(picked.get("Title",""))
            row["OutcomeLabel"] = out_label
            row["Effect_measure"] = eff_measure
            row["Effect"] = eff
            row["Lower_CI"] = lcl
            row["Upper_CI"] = ucl
            row["Notes"] = notes
            df_ex = st.session_state.get("extract_df")
            df_ex = ensure_columns(df_ex, schema_cols, default="")
            st.session_state["extract_df"] = pd.concat([df_ex, pd.DataFrame([row])], ignore_index=True)
            st.success("已新增一筆到 extraction 寬表（可在下方 editor 一次修改多欄）。")
            st.session_state["extract_saved"] = False

        st.markdown("---")
        st.markdown("#### " + t("extract_editor"))
        st.caption("建議：先在此一次編輯整張表，按『儲存/commit』後再到 Step 6 跑 MA/森林圖。")

        df_ex = st.session_state.get("extract_df")
        df_ex = ensure_columns(df_ex, schema_cols, default="")
        edited = st.data_editor(
            df_ex,
            use_container_width=True,
            num_rows="dynamic",
            key="extract_editor_table",
        )

        cA, cB = st.columns([1,2])
        with cA:
            if st.button(t("extract_save"), type="primary"):
                # Commit edited table
                st.session_state["extract_df"] = edited.copy()
                st.session_state["extract_saved"] = True
                st.success("已儲存/commit。接下來 Step 6 會使用這張表。")
        with cB:
            st.download_button("Download extraction CSV", data=to_csv_bytes(edited), file_name="extraction_wide.csv", mime="text/csv")

        # Validation preview (non-blocking, red warnings)
        st.markdown("---")
        st.markdown("#### Validation (non-blocking)")
        out_filter = st.text_input( t("ma_outcome_label"), value=st.session_state.get("ma_outcome_input",""), key="ma_outcome_input_step5", help="例如：visual acuity / photic / defocus；用 substring 匹配。")
        _meas_opts = ["OR","RR","HR","MD","SMD"]
        _meas_default = st.session_state.get("ma_measure_choice","OR")
        _meas_idx = _meas_opts.index(_meas_default) if _meas_default in _meas_opts else 0
        meas = st.selectbox(t("ma_measure"), options=_meas_opts, index=_meas_idx, key="ma_measure_choice_step5")
        # Sync into downstream defaults (Step 6)
        st.session_state["ma_outcome_input"] = out_filter
        st.session_state["ma_measure_choice"] = meas
        rep = validate_extraction(edited, meas, out_filter)
        if rep.empty:
            st.caption("No rows matched the current outcome/measure filter, or extraction table is empty.")
        else:
            n_bad = int((~rep["ok"]).sum())
            if n_bad > 0:
                st.error(f"Rows with issues: {n_bad} (still allowed to proceed; Step 6 will skip invalid rows).")
            st.dataframe(rep, use_container_width=True)

# =========================================================
# Tab 6: Step 6 MA + Forest (button run; stable UI)
# =========================================================

with tabs[6]:
    st.subheader(t("tabs_ma"))
    df_ex = st.session_state.get("extract_df")
    if df_ex is None or not isinstance(df_ex, pd.DataFrame) or df_ex.empty:
        st.info("尚無 extraction 寬表。請先在 Step 5 建立/儲存。")
    else:
        df_ex = df_ex.copy()
        df_ex = ensure_columns(df_ex, ["OutcomeLabel","Effect_measure","Effect","SE","Lower_CI","Upper_CI","First_author","Year","Title","PMID"], default="")

        outcome = st.text_input(t("ma_outcome_label"), key="ma_outcome_input", help="用 substring 匹配 OutcomeLabel。留空=全選。")
        measure = st.selectbox(t("ma_measure"), options=["OR","RR","HR","MD","SMD"], key="ma_measure_choice")
        model = st.selectbox("Model", options=["Fixed effect", "Random effects (DL)"], key="ma_model_choice")

        st.caption("按下 Run 之後才會執行（避免你在 Step 5 編輯時反覆 rerun）。")

        if st.button(t("ma_run"), type="primary"):
            work = df_ex.copy()
            if outcome:
                work = work[work["OutcomeLabel"].astype(str).str.lower().str.contains(outcome.lower(), na=False)]
            work = work[work["Effect_measure"].astype(str).str.upper().str.contains(measure.upper(), na=False)]
            if work.empty:
                st.error("沒有符合 OutcomeLabel/Effect_measure 的列。")
            else:
                if model.startswith("Random"):
                    res, skipped = random_effect_meta(work, measure.upper().strip())
                else:
                    res, skipped = fixed_effect_meta(work, measure.upper().strip())

                st.session_state["ma_last_result"] = res
                st.session_state["ma_skipped_rows"] = skipped

                if res is None:
                    st.error("沒有可用的列（可能 CI/SE/Effect 不合法）。請看下方 skipped rows。")
                else:
                    pooled = res["pooled"]; lcl = res["pooled_lcl"]; ucl = res["pooled_ucl"]
                    st.success(f"{res['model']} 完成：k={res['k']}｜pooled={pooled:.3g} (95% CI {lcl:.3g}–{ucl:.3g})｜I²={res.get('I2',0):.1f}%")
                    st.markdown("##### Forest plot（RevMan-like）")
                    plot_forest(res, title=f"{outcome or 'Outcome'} ({measure}, {res['model']})")
                    st.markdown("##### Included rows")
                    st.dataframe(res["study_table"], use_container_width=True)

        skipped = st.session_state.get("ma_skipped_rows")
        if isinstance(skipped, pd.DataFrame) and not skipped.empty:
            st.markdown("##### Skipped rows (with reasons)")
            cols = [c for c in ["PMID","First_author","Year","OutcomeLabel","Effect_measure","Effect","Lower_CI","Upper_CI","SE","SkipReason"] if c in skipped.columns]
            st.dataframe(skipped[cols], use_container_width=True)

with tabs[7]:
    st.subheader(t("tabs_rob2"))
    df = st.session_state.get("pubmed_records")
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        st.info("請先 Run 並抓到 records。")
    else:
        include_pmids = [pmid for pmid, dec in (st.session_state.get("ft_decision") or {}).items() if dec == "Include for meta-analysis"]
        if not include_pmids:
            st.warning("目前沒有 Full-text decision = Include for meta-analysis 的研究；ROB 2.0 通常在納入後做。你可先建立空白評分，或回 Step 4b 完成全文決策。")

        with st.expander("ROB 2.0 traffic light（依目前已儲存的評分）", expanded=True):
            rob_map = st.session_state.get("rob2", {}) or {}
            pmids_plot = [str(p) for p in include_pmids[:30]] if include_pmids else []
            if pmids_plot:
                plot_rob2_traffic_light(rob_map, pmids_plot, title="RoB 2.0 traffic light (top 30)")
            else:
                st.info("尚未有納入研究可畫圖。")


        with st.expander("ROB 2.0 traffic light（依目前已儲存的評分）", expanded=True):
            rob_map = st.session_state.get("rob2", {}) or {}
            pmids_plot = [str(p) for p in include_pmids[:30]] if include_pmids else []
            if pmids_plot:
                plot_rob2_traffic_light(rob_map, pmids_plot, title="RoB 2.0 traffic light (top 30)")
            else:
                st.info("尚未有納入研究可畫圖。")


        df = ensure_columns(df.copy(), ["PMID","Title","First_author","Year"], default="")
        for pmid in include_pmids[:50]:
            rr = df[df["PMID"].astype(str) == str(pmid)]
            title = rr["Title"].iloc[0] if not rr.empty else ""
            fa = rr["First_author"].iloc[0] if not rr.empty else ""
            year = rr["Year"].iloc[0] if not rr.empty else ""

            st.markdown(f"### {fa} ({year}) — {short(title, 110)}")
            rob = st.session_state["rob2"].get(pmid, {})
            cols = st.columns([1,1,1,1,1])
            for i, dom in enumerate(ROB_DOMAINS):
                with cols[i]:
                    rob[dom] = st.selectbox(dom, options=ROB_LEVELS, index=ROB_LEVELS.index(rob.get(dom,"Unclear")) if rob.get(dom,"Unclear") in ROB_LEVELS else 3, key=f"rob_{pmid}_{i}")
            rob["Overall"] = st.selectbox("Overall Risk of Bias", options=ROB_LEVELS, index=ROB_LEVELS.index(rob.get("Overall","Unclear")) if rob.get("Overall","Unclear") in ROB_LEVELS else 3, key=f"rob_{pmid}_overall")
            rob["Rationale"] = st.text_area("Rationale / notes (required for transparency)", value=rob.get("Rationale",""), key=f"rob_{pmid}_why", height=90)
            st.session_state["rob2"][pmid] = rob

            if llm_available():
                with st.expander("BYOK：AI 建議 ROB 2.0（需全文文字；請人工核對）", expanded=False):
                    ft_text = st.session_state.get("ft_text", {}).get(pmid, "")
                    if not ft_text.strip():
                        st.info("此篇尚無 full-text text。請先在 Step 4b 上傳 PDF 抽字或貼上全文（或貼上 Methods/Results）。")
                    else:
                        st.session_state.setdefault("rob2_ai", {})
                        existing = st.session_state["rob2_ai"].get(pmid)

                        if existing:
                            st.caption("已存在 AI 建議（如下）。你也可以按下方按鈕重新產生。")
                            st.json(existing)

                        if st.button("Generate / Refresh ROB2 suggestion", key=f"rob2_ai_{pmid}"):
                            sys = (
                                "You are an evidence synthesis methodologist. "
                                "Assess Cochrane RoB 2.0 for an RCT based on the provided full text. "
                                "Return STRICT JSON only (no markdown). "
                                "Required keys: "
                                "Randomization process, Deviations from intended interventions, Missing outcome data, "
                                "Measurement of the outcome, Selection of the reported result, Overall, Rationale. "
                                "Each domain value must be exactly one of: Low, Some concerns, High, Unclear. "
                                "Rationale should be short, cite concrete signals from the text."
                            )
                            user = json.dumps(
                                {
                                    "pmid": pmid,
                                    "title": title,
                                    "year": year,
                                    "journal": journal,
                                    "full_text": ft_text[:20000],
                                },
                                ensure_ascii=False,
                            )
                            d = llm_json(sys, user, max_tokens=1400)
                            if not d:
                                st.warning("AI 未回傳可解析的 JSON（或未啟用 LLM）。請確認 sidebar 已填入可用的 API key / Base URL / Model。")
                            else:
                                st.session_state["rob2_ai"][pmid] = d
                                # Auto-fill the manual ROB form (will take effect on next rerun)
                                rob_now = (st.session_state.get("rob2", {}) or {}).get(pmid, {}) or {}
                                for dom in ROB_DOMAINS:
                                    v = d.get(dom)
                                    if isinstance(v, str) and v in ROB_LEVELS:
                                        rob_now[dom] = v
                                ov = d.get("Overall")
                                if isinstance(ov, str) and ov in ROB_LEVELS:
                                    rob_now["Overall"] = ov
                                rat = d.get("Rationale")
                                if isinstance(rat, str) and rat.strip():
                                    rob_now["Rationale"] = rat.strip()
                                st.session_state.setdefault("rob2", {})
                                st.session_state["rob2"][pmid] = rob_now
                                st.success("已將 AI 建議寫入此篇 ROB 2.0 表單（你仍可在上方手動調整）。")
                                st.json(d)


# =========================================================
# Tab 8: Manuscript (sections shown + optional BYOK) + export
# =========================================================
with tabs[8]:
    st.subheader(t("tabs_ms"))
    proto: Protocol = st.session_state.get("protocol")
    df = st.session_state.get("pubmed_records")
    prisma = compute_prisma(df)
    ma_res = st.session_state.get("ma_last_result")

    if not proto:
        st.info("請先 Run。")
    else:
        st.caption("此處會顯示『分段』稿件草稿；若缺資訊，用『』標示讓研究者後續補填。")
        style_notes = st.text_area("Writing style notes (optional; for BYOK)", key="writing_style_notes", height=80,
                                   help="可貼你的範本風格要點（例如 AJO/BJO 語氣、段落結構）。不貼也可。")

        # Basic draft always available
        basic = generate_manuscript_basic(proto, prisma, ma_res)
        st.session_state["ms_sections"] = basic

        # Optional BYOK enhancement
        if llm_available():
            if st.button(t("ms_generate"), type="primary"):
                d2 = manuscript_llm_enhance(proto, prisma, ma_res, style_notes) or basic
                st.session_state["ms_sections"] = d2

        ms = st.session_state.get("ms_sections") or basic
        for sec in ["Introduction","Methods","Results","Discussion"]:
            with st.expander(sec, expanded=True):
                st.write(ms.get(sec,""))

        # Export DOCX
        st.markdown("---")
        if not HAS_DOCX:
            st.warning("環境未安裝 python-docx：無法匯出 Word。你可直接複製上方文字到 Word。")
        else:
            if st.button(t("export_docx")):
                doc = Document()
                style = doc.styles["Normal"]
                style.font.name = "Times New Roman"
                style.font.size = Pt(11)

                doc.add_heading(st.session_state.get("question_en") or st.session_state.get("question") or "Meta-analysis draft", level=1)
                doc.add_paragraph("Disclaimer: For academic use only. Verify all results and citations.")

                doc.add_heading("Protocol", level=2)
                doc.add_paragraph(json.dumps(proto.to_dict(), ensure_ascii=False, indent=2))

                doc.add_heading("PRISMA counts", level=2)
                doc.add_paragraph(json.dumps(prisma, ensure_ascii=False, indent=2))

                if ma_res:
                    doc.add_heading("Meta-analysis (fixed effect)", level=2)
                    doc.add_paragraph(f"{ma_res['measure']} pooled = {ma_res['pooled']:.4g} (95% CI {ma_res['pooled_lcl']:.4g} to {ma_res['pooled_ucl']:.4g}); k={ma_res['k']}")

                for sec in ["Introduction","Methods","Results","Discussion"]:
                    doc.add_heading(sec, level=2)
                    doc.add_paragraph(ms.get(sec,""))

                # Save
                out = io.BytesIO()
                doc.save(out)
                out.seek(0)
                st.download_button("Download DOCX", data=out.getvalue(), file_name="meta_analysis_draft.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# =========================================================
# Tab 9: Diagnostics
# =========================================================
with tabs[9]:
    st.subheader(t("tabs_diag"))
    diag = st.session_state.get("diagnostics") or {}
    st.caption("當 PubMed 被擋（回 HTML / 403 / 連線失敗）時，這裡最重要。")
    st.json(diag)
    if diag.get("esearch_urls"):
        st.write("esearch_urls:")
        st.code("\n".join(diag.get("esearch_urls", [])), language="text")
    elif diag.get("esearch_url"):
        st.write("esearch_url:")
        st.code(diag["esearch_url"], language="text")
    if diag.get("efetch_urls"):
        st.write("efetch_urls:")
        st.code("\n".join(diag["efetch_urls"]), language="text")
